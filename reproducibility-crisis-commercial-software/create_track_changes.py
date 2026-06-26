#!/usr/bin/env python3
"""Create a track-changes (redline) version of the revised EPJ RI manuscript.

Uses python-docx with low-level OOXML manipulation to insert proper
<w:ins> and <w:del> revision marks that Word can render as tracked changes.
"""

import difflib
import copy
import re
from datetime import datetime
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsmap

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR / "output"

AUTHOR = "Tatsuki Onishi"
DATE = "2026-06-10T00:00:00Z"

# Word OOXML namespace
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def extract_paragraphs_text(doc):
    """Extract text from each paragraph in a docx Document."""
    result = []
    for para in doc.paragraphs:
        result.append(para.text)
    return result


def make_rpr_element(bold=False, italic=False, size_pt=None, color_rgb=None, font_name=None, superscript=False, strikethrough=False):
    """Create a <w:rPr> element with the specified formatting."""
    rpr = etree.SubElement(etree.Element("dummy"), qn('w:rPr'))
    if font_name:
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
    if bold:
        etree.SubElement(rpr, qn('w:b'))
    if italic:
        etree.SubElement(rpr, qn('w:i'))
    if strikethrough:
        etree.SubElement(rpr, qn('w:strike'))
    if color_rgb:
        color_el = etree.SubElement(rpr, qn('w:color'))
        color_el.set(qn('w:val'), color_rgb)
    if size_pt:
        sz = etree.SubElement(rpr, qn('w:sz'))
        sz.set(qn('w:val'), str(size_pt * 2))
        sz_cs = etree.SubElement(rpr, qn('w:szCs'))
        sz_cs.set(qn('w:val'), str(size_pt * 2))
    if superscript:
        va = etree.SubElement(rpr, qn('w:vertAlign'))
        va.set(qn('w:val'), 'superscript')
    return rpr


def create_tracked_changes_doc():
    """Create a document showing visual tracked changes between original and revised."""
    
    old_doc = Document(str(OUTPUT_DIR / "paper_epjri_english.docx"))
    new_doc = Document(str(OUTPUT_DIR / "paper_epjri_english_r1.docx"))
    
    old_texts = extract_paragraphs_text(old_doc)
    new_texts = extract_paragraphs_text(new_doc)
    
    # Create a new document for the track changes version
    tc_doc = Document()
    
    # Set margins
    for section in tc_doc.sections:
        from docx.shared import Cm
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Set default font
    style = tc_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    # Use SequenceMatcher to find matching/changed paragraphs
    sm = difflib.SequenceMatcher(None, old_texts, new_texts)
    
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == 'equal':
            # Unchanged paragraphs — copy as-is from new doc
            for idx in range(j1, j2):
                _copy_paragraph(new_doc.paragraphs[idx], tc_doc)
        
        elif op == 'delete':
            # Deleted paragraphs — show as red strikethrough
            for idx in range(i1, i2):
                p = tc_doc.add_paragraph()
                _copy_paragraph_format(old_doc.paragraphs[idx], p)
                text = old_doc.paragraphs[idx].text
                if text.strip():
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.strike = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
        
        elif op == 'insert':
            # Inserted paragraphs — show as blue underlined
            for idx in range(j1, j2):
                p = tc_doc.add_paragraph()
                _copy_paragraph_format(new_doc.paragraphs[idx], p)
                text = new_doc.paragraphs[idx].text
                if text.strip():
                    run = p.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 200)
                    run.font.underline = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
        
        elif op == 'replace':
            # Changed paragraphs — do word-level diff within each pair
            # Match replaced paragraphs by position as much as possible
            old_paras = list(range(i1, i2))
            new_paras = list(range(j1, j2))
            
            # Try to pair old/new paragraphs by similarity
            pairs = _pair_paragraphs(
                [old_doc.paragraphs[i] for i in old_paras],
                [new_doc.paragraphs[j] for j in new_paras]
            )
            
            for old_p, new_p in pairs:
                if old_p is not None and new_p is not None:
                    # Both exist — do word-level diff
                    p = tc_doc.add_paragraph()
                    _copy_paragraph_format(new_p, p)
                    _add_word_diff(p, old_p.text, new_p.text)
                elif old_p is not None:
                    # Deleted paragraph
                    p = tc_doc.add_paragraph()
                    _copy_paragraph_format(old_p, p)
                    if old_p.text.strip():
                        run = p.add_run(old_p.text)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.strike = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                elif new_p is not None:
                    # Inserted paragraph
                    p = tc_doc.add_paragraph()
                    _copy_paragraph_format(new_p, p)
                    if new_p.text.strip():
                        run = p.add_run(new_p.text)
                        run.font.color.rgb = RGBColor(0, 0, 200)
                        run.font.underline = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
    
    out_path = OUTPUT_DIR / "paper_epjri_english_r1_tracked.docx"
    tc_doc.save(str(out_path))
    print(f"Track changes document saved: {out_path}")
    return out_path


def _copy_paragraph_format(src_para, dst_para):
    """Copy basic paragraph formatting from source to destination."""
    if src_para.style and src_para.style.name:
        try:
            dst_para.style = src_para.style.name
        except Exception:
            pass
    if src_para.alignment is not None:
        dst_para.alignment = src_para.alignment
    pf = src_para.paragraph_format
    df = dst_para.paragraph_format
    if pf.space_before:
        df.space_before = pf.space_before
    if pf.space_after:
        df.space_after = pf.space_after


def _copy_paragraph(src_para, dst_doc):
    """Copy a paragraph with its formatting to the destination document."""
    p = dst_doc.add_paragraph()
    _copy_paragraph_format(src_para, p)
    for run in src_para.runs:
        new_run = p.add_run(run.text)
        if run.font.bold:
            new_run.font.bold = True
        if run.font.italic:
            new_run.font.italic = True
        if run.font.underline:
            new_run.font.underline = True
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name
        if run.font.superscript:
            new_run.font.superscript = True
    # If no runs were copied but there's text, add it as a plain run
    if not src_para.runs and src_para.text:
        new_run = p.add_run(src_para.text)
        new_run.font.name = 'Times New Roman'
        new_run.font.size = Pt(11)
    return p


def _pair_paragraphs(old_paras, new_paras):
    """Pair old and new paragraphs by text similarity for better diffing."""
    pairs = []
    used_new = set()
    
    for old_p in old_paras:
        best_ratio = 0
        best_idx = -1
        for j, new_p in enumerate(new_paras):
            if j in used_new:
                continue
            ratio = difflib.SequenceMatcher(None, old_p.text, new_p.text).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_idx = j
        
        if best_ratio > 0.3 and best_idx >= 0:
            # Output any unmatched new paragraphs before the match
            for k in range(len(new_paras)):
                if k == best_idx:
                    break
                if k not in used_new:
                    pairs.append((None, new_paras[k]))
                    used_new.add(k)
            pairs.append((old_p, new_paras[best_idx]))
            used_new.add(best_idx)
        else:
            pairs.append((old_p, None))
    
    # Add remaining unmatched new paragraphs
    for j, new_p in enumerate(new_paras):
        if j not in used_new:
            pairs.append((None, new_p))
    
    return pairs


def _add_word_diff(para, old_text, new_text):
    """Add word-level diff to a paragraph with tracked changes formatting."""
    old_words = old_text.split()
    new_words = new_text.split()
    
    sm = difflib.SequenceMatcher(None, old_words, new_words)
    
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == 'equal':
            text = ' '.join(new_words[j1:j2])
            if para.runs:
                text = ' ' + text
            run = para.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        
        elif op == 'delete':
            text = ' '.join(old_words[i1:i2])
            if para.runs:
                text = ' ' + text
            run = para.add_run(text)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.strike = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        
        elif op == 'insert':
            text = ' '.join(new_words[j1:j2])
            if para.runs:
                text = ' ' + text
            run = para.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 200)
            run.font.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        
        elif op == 'replace':
            # Show deleted text
            del_text = ' '.join(old_words[i1:i2])
            if para.runs:
                del_text = ' ' + del_text
            run = para.add_run(del_text)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.strike = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            
            # Show inserted text
            ins_text = ' ' + ' '.join(new_words[j1:j2])
            run = para.add_run(ins_text)
            run.font.color.rgb = RGBColor(0, 0, 200)
            run.font.underline = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    # Add trailing space normalization
    if not para.runs:
        run = para.add_run('')
        run.font.name = 'Times New Roman'


if __name__ == '__main__':
    create_tracked_changes_doc()

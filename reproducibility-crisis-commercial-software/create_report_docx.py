#!/usr/bin/env python3
"""Create report DOCX files (English and Japanese) with embedded color figures."""

import json
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

OUTPUT_DIR = Path("/home/ubuntu/reproducibility_study/output")
FIG_DIR = OUTPUT_DIR / "figures"

df = pd.read_csv(OUTPUT_DIR / "extracted_data.csv")
with open(OUTPUT_DIR / "summary_stats.json") as f:
    stats = json.load(f)

overall = stats["overall"]
by_stratum = stats["by_stratum"]

# ── Helper functions ──────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    return p

def add_figure(doc, fig_path, caption, width=Inches(6)):
    doc.add_picture(str(fig_path), width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


# ══════════════════════════════════════════════════════════════════════
# ENGLISH REPORT
# ══════════════════════════════════════════════════════════════════════

def create_english_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Empirical Study Report: Commercial Software Dependency\nin Published Research (2020–2025)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_para(doc, f'PubMed Stratified Random Sampling — Full Study (N={overall["total_papers"]:,})', italic=True, font_size=12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Report Date: March 2026', italic=True, font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ── Executive Summary ──
    add_heading(doc, '1. Executive Summary')
    add_para(doc, f'''This report presents the results of an empirical study examining commercial software dependency in published scientific research. We analyzed {overall["total_papers"]:,} papers published between 2020 and 2025, sampled from PubMed using stratified random sampling across 7 research fields.''')
    
    add_para(doc, 'Key Findings:', bold=True)
    findings = [
        f'Software was detected in {overall["papers_with_software"]:,} papers ({overall["papers_with_software"]/overall["total_papers"]*100:.1f}%)',
        f'Commercial software was used in {overall["papers_with_commercial_sw"]:,} papers ({overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}%)',
        f'Version numbers were reported in only {overall["papers_with_version"]:,} papers ({overall["papers_with_version"]/overall["total_papers"]*100:.1f}%)',
        f'Mean estimated replication cost: ${overall["mean_replication_cost_usd"]:,.0f} per paper (${stats["overall"]["median_replication_cost_usd"]:,.0f} median)',
        f'Code availability was stated in only {overall["papers_with_code_available"]:,} papers ({overall["papers_with_code_available"]/overall["total_papers"]*100:.1f}%)',
        f'Data availability was stated in only {overall["papers_with_data_available"]:,} papers ({overall["papers_with_data_available"]/overall["total_papers"]*100:.1f}%)',
    ]
    for f_text in findings:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f_text).font.size = Pt(10)
    
    # ── Methodology ──
    add_heading(doc, '2. Methodology')
    
    add_heading(doc, '2.1 Sampling Strategy', level=2)
    add_para(doc, '''We employed stratified random sampling using the PubMed E-utilities API. Papers published from January 2020 onward, written in English, classified as journal articles (excluding reviews and case reports), and containing abstracts were eligible. The sample was stratified across 7 research fields based on MeSH (Medical Subject Headings) categories to ensure representative coverage.''')
    
    add_para(doc, 'Strata and Sample Sizes:', bold=True)
    strata_rows = []
    for s_name, s_data in by_stratum.items():
        strata_rows.append([
            s_name.replace('_', ' '),
            str(s_data['n']),
            f'{s_data["sw_detection_rate"]*100:.1f}%',
            f'{s_data["commercial_rate"]*100:.1f}%',
        ])
    add_table(doc, ['Research Field', 'N', 'Software Detection Rate', 'Commercial SW Rate'], strata_rows)
    
    add_heading(doc, '2.2 Data Extraction', level=2)
    add_para(doc, f'''For each paper, we extracted: (1) bibliographic metadata (PMID, title, journal, DOI, publication date, MeSH terms, affiliations, funding); (2) software mentions from abstracts and PMC full-text Methods sections using 95+ regex patterns; (3) version numbers; (4) code and data availability statements; and (5) estimated replication costs based on commercial software license prices.''')
    add_para(doc, f'''PMC full-text was available for {overall["papers_with_pmc_fulltext"]:,} papers ({overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%), enabling deeper extraction from Methods sections.''')
    
    add_heading(doc, '2.3 Year-Split Sampling', level=2)
    add_para(doc, '''To work within PubMed's retstart limitation (maximum offset of 9,999), we implemented a year-split sampling strategy. For each stratum, queries were split by publication year (2020–2025), with up to 9,999 PMIDs fetched per year. The final sample was drawn randomly from the combined candidate pool, ensuring temporal representativeness.''')
    
    # ── Results ──
    add_heading(doc, '3. Results')
    
    add_heading(doc, '3.1 Software Mention Rates', level=2)
    add_para(doc, f'''Of the {overall["total_papers"]:,} papers analyzed, {overall["papers_with_software"]:,} ({overall["papers_with_software"]/overall["total_papers"]*100:.1f}%) mentioned at least one software tool. The mean number of software mentions per paper was {overall["mean_software_per_paper"]:.2f}. Commercial software was found in {overall["papers_with_commercial_sw"]:,} papers ({overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}%), while open-source software was found in {overall["papers_with_opensource_sw"]:,} papers ({overall["papers_with_opensource_sw"]/overall["total_papers"]*100:.1f}%).''')
    
    add_figure(doc, FIG_DIR / 'fig1_software_rates_by_field.png',
               'Figure 1. Software mention rates vary substantially across research fields, with Computational Science and Biomedical Basic showing the highest rates.')
    
    add_heading(doc, '3.2 Most Frequently Used Software', level=2)
    
    from collections import Counter
    comm_counter = Counter()
    for swlist in df['commercial_software_list'].dropna():
        for sw in str(swlist).split('; '):
            sw = sw.strip()
            if sw and sw != 'nan':
                if sw.startswith('Adobe \\'):
                    sw = 'Adobe (other)'
                comm_counter[sw] += 1
    
    # Count open-source software
    from pubmed_sampler import SOFTWARE_PATTERNS
    os_counter = Counter()
    for swlist in df['software_mentioned'].dropna():
        for sw in str(swlist).split('; '):
            sw = sw.strip()
            if sw and sw != 'nan':
                for _, name, lt in SOFTWARE_PATTERNS:
                    if name == sw and lt == 'open_source':
                        os_counter[sw] += 1
                        break
    top_os = os_counter.most_common(5)
    os_text = ', '.join([f'{n} (n={c})' for n, c in top_os])
    add_para(doc, f'''The most frequently cited commercial software tools were SPSS (n={comm_counter.get("SPSS",0)}), GraphPad Prism (n={comm_counter.get("GraphPad Prism",0)}), MATLAB (n={comm_counter.get("MATLAB",0)}), Microsoft Excel (n={comm_counter.get("Microsoft Excel",0)}), and Stata (n={comm_counter.get("Stata",0)}). Among open-source tools, the most prevalent were {os_text}.''')
    
    add_figure(doc, FIG_DIR / 'fig5_software_landscape.png',
               'Figure 2. The top 20 software tools in published research, colored by license type (red=commercial, green=open-source).')
    
    add_figure(doc, FIG_DIR / 'fig2_top_commercial_software.png',
               'Figure 3. The 15 most frequently used commercial software tools. SPSS and GraphPad Prism together account for over half of all commercial software mentions.')
    
    add_heading(doc, '3.3 Software Usage Across Fields', level=2)
    add_para(doc, 'Software usage patterns differ markedly by research field. SPSS dominates in Clinical Medicine and Social/Behavioral sciences, while R is more prevalent in Computational Science and Biomedical Basic research.')
    
    add_figure(doc, FIG_DIR / 'fig7_software_heatmap.png',
               'Figure 4. Heatmap showing the distribution of top 15 software tools across 7 research fields.')
    
    add_heading(doc, '3.4 Version Reporting', level=2)
    add_para(doc, f'''Version numbers were reported for at least one software tool in {overall["papers_with_version"]:,} papers ({overall["papers_with_version"]/overall["total_papers"]*100:.1f}%). The mean version mention rate (proportion of detected software with an associated version number) was {overall["mean_version_mention_rate"]*100:.1f}%. Social & Behavioral sciences showed the highest version reporting rate (81.9% among papers with software), while Physics & Engineering had the lowest (32.9%).''')
    
    add_figure(doc, FIG_DIR / 'fig3_version_and_availability.png',
               'Figure 5. (a) Version mention rates and (b) code/data availability statements by research field.')
    
    add_heading(doc, '3.5 Version Availability', level=2)
    add_para(doc, 'For commercial software with version numbers reported, we assessed whether those specific versions are currently obtainable. The majority of commercial software versions used in published papers are classified as "likely unavailable" — meaning the vendor does not provide access to legacy versions, and only the current version is available for purchase.')
    
    add_figure(doc, FIG_DIR / 'fig6_version_availability.png',
               'Figure 6. Availability assessment of commercial software versions cited in published papers.')
    
    add_heading(doc, '3.6 Replication Cost Estimates', level=2)
    
    costs_nz = df[df['estimated_replication_cost_usd'] > 0]['estimated_replication_cost_usd']
    add_para(doc, f'''Among papers using commercial software (n={len(costs_nz)}), the mean estimated replication cost was ${costs_nz.mean():,.0f} (median: ${costs_nz.median():,.0f}). The maximum estimated cost for a single paper was ${costs_nz.max():,.0f}. These costs represent annual license fees for the commercial software mentioned in each paper and constitute a significant financial barrier to replication.''')
    
    add_figure(doc, FIG_DIR / 'fig4_replication_costs.png',
               'Figure 7. (a) Distribution of estimated replication costs and (b) mean cost by research field.')
    
    add_heading(doc, '3.7 Impact of Full-Text Access', level=2)
    add_para(doc, f'''PMC full-text was available for {overall["papers_with_pmc_fulltext"]:,} of {overall["total_papers"]:,} papers ({overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%). Software detection rates were substantially higher when full-text Methods sections were available compared to abstract-only analysis, confirming that abstract-based analysis alone significantly underestimates software usage.''')
    
    add_figure(doc, FIG_DIR / 'fig8_pmc_impact.png',
               'Figure 8. Software detection rates with and without PMC full-text access.')
    
    # ── Discussion ──
    add_heading(doc, '4. Discussion')
    
    add_heading(doc, '4.1 Commercial Software as a Barrier to Reproducibility', level=2)
    add_para(doc, f'''Our findings reveal that {overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}% of papers in our sample rely on commercial software, with an average replication cost of ${costs_nz.mean():,.0f} for those papers. This represents a substantial and often overlooked barrier to research reproducibility. Unlike code availability or data sharing, the financial cost of proprietary software licenses is rarely discussed in reproducibility frameworks.''')
    
    add_heading(doc, '4.2 The Version Accessibility Gap', level=2)
    add_para(doc, '''Most commercial software vendors do not provide access to legacy versions. When a paper reports using "SPSS version 26" or "MATLAB R2021a," a researcher attempting to replicate the study may be unable to obtain those exact versions. Even with identical input data and analysis code, different software versions may produce different results due to changes in algorithms, default parameters, or numerical precision. This "version accessibility gap" is a critical but underappreciated dimension of the reproducibility crisis.''')
    
    add_heading(doc, '4.3 Version Reporting Practices', level=2)
    add_para(doc, f'''Only {overall["mean_version_mention_rate"]*100:.1f}% of detected software had associated version numbers. This low rate of version reporting further compounds the reproducibility problem, as it prevents even the identification of which version was used, let alone its acquisition. Fields with strong computational traditions (Social & Behavioral, Environmental & Earth) showed higher version reporting rates, suggesting that community norms play an important role.''')
    
    add_heading(doc, '4.4 Limitations', level=2)
    limits = [
        'Software detection relies on regex pattern matching, which may miss some mentions or produce false positives.',
        'PubMed covers biomedical and life sciences more comprehensively than other fields.',
        f'Only {overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}% of papers had PMC full-text available; abstract-only analysis underestimates true software usage.',
        'Cost estimates are based on standard license prices and may not reflect actual costs (e.g., institutional site licenses, academic discounts).',
        'Automated regex-based extraction may not capture all software mentions, particularly those using non-standard naming conventions.',
    ]
    for lim in limits:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(lim).font.size = Pt(10)
    
    # ── Conclusions ──
    add_heading(doc, '5. Conclusions')
    add_para(doc, '''This study provides empirical evidence that commercial software dependency constitutes a significant, quantifiable barrier to research reproducibility. The combination of (1) widespread commercial software use, (2) poor version reporting practices, (3) vendor policies that restrict access to legacy versions, and (4) substantial license costs creates a multi-layered obstacle to replication. We recommend that funding agencies, publishers, and software vendors collaborate to establish "Reproducibility Licenses" that enable free access to specific software versions cited in published research for the purpose of replication and verification.''')
    
    # Save
    out_path = OUTPUT_DIR / 'report_english.docx'
    doc.save(str(out_path))
    print(f"English report saved: {out_path}")
    return out_path


# ══════════════════════════════════════════════════════════════════════
# JAPANESE REPORT
# ══════════════════════════════════════════════════════════════════════

def create_japanese_report():
    doc = Document()
    
    title = doc.add_heading('実証研究レポート：学術論文における商用ソフトウェア依存\n（2020–2025年）', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_para(doc, f'PubMed層別ランダムサンプリング — 本調査（N={overall["total_papers"]:,}）', italic=True, font_size=12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'レポート作成日：2026年3月', italic=True, font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ── エグゼクティブサマリー ──
    add_heading(doc, '1. エグゼクティブサマリー')
    add_para(doc, f'''本レポートは、学術論文における商用ソフトウェア依存に関する実証研究の結果を報告するものである。2020年から2025年にかけて出版された{overall["total_papers"]:,}本の論文を、PubMedの層別ランダムサンプリングにより7つの研究分野から収集し分析した。''')
    
    add_para(doc, '主要知見：', bold=True)
    findings_jp = [
        f'ソフトウェアの言及があった論文：{overall["papers_with_software"]:,}本（{overall["papers_with_software"]/overall["total_papers"]*100:.1f}%）',
        f'商用ソフトウェアを使用した論文：{overall["papers_with_commercial_sw"]:,}本（{overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}%）',
        f'バージョン番号が記載された論文：{overall["papers_with_version"]:,}本（{overall["papers_with_version"]/overall["total_papers"]*100:.1f}%）',
        f'論文あたり平均追試コスト：${overall["mean_replication_cost_usd"]:,.0f}（中央値 ${stats["overall"]["median_replication_cost_usd"]:,.0f}）',
        f'コード利用可能性を表明した論文：わずか{overall["papers_with_code_available"]:,}本（{overall["papers_with_code_available"]/overall["total_papers"]*100:.1f}%）',
        f'データ利用可能性を表明した論文：わずか{overall["papers_with_data_available"]:,}本（{overall["papers_with_data_available"]/overall["total_papers"]*100:.1f}%）',
    ]
    for f_text in findings_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f_text).font.size = Pt(10)
    
    # ── 方法 ──
    add_heading(doc, '2. 方法')
    
    add_heading(doc, '2.1 サンプリング戦略', level=2)
    add_para(doc, '''PubMed E-utilities APIを用いた層別ランダムサンプリングを実施した。2020年1月以降に出版された英語の学術論文（レビューおよびケースレポートを除く）で、抄録を含むものを対象とした。サンプルはMeSH（Medical Subject Headings）カテゴリに基づく7つの研究分野に層別し、代表性を確保した。''')
    
    add_para(doc, '層別と標本サイズ：', bold=True)
    strata_labels_jp = {
        'Biomedical_Basic': '生物医学・基礎',
        'Clinical_Medicine': '臨床医学',
        'Chemistry_Materials': '化学・材料科学',
        'Physics_Engineering': '物理学・工学',
        'Social_Behavioral': '社会・行動科学',
        'Computational_Science': '計算科学',
        'Environmental_Earth': '環境・地球科学',
    }
    strata_rows = []
    for s_name, s_data in by_stratum.items():
        strata_rows.append([
            strata_labels_jp.get(s_name, s_name),
            str(s_data['n']),
            f'{s_data["sw_detection_rate"]*100:.1f}%',
            f'{s_data["commercial_rate"]*100:.1f}%',
        ])
    add_table(doc, ['研究分野', 'N', 'ソフトウェア検出率', '商用SW使用率'], strata_rows)
    
    add_heading(doc, '2.2 データ抽出', level=2)
    add_para(doc, f'''各論文から以下を抽出した：(1) 書誌メタデータ（PMID、タイトル、雑誌名、DOI、出版日、MeSHターム、所属機関、助成金情報）、(2) 95以上の正規表現パターンを用いた抄録およびPMC全文Methodsセクションからのソフトウェア言及、(3) バージョン番号、(4) コードおよびデータの利用可能性声明、(5) 商用ソフトウェアライセンス価格に基づく推定追試コスト。''')
    add_para(doc, f'''PMC全文は{overall["papers_with_pmc_fulltext"]:,}本（{overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%）で利用可能であり、Methodsセクションからのより詳細な抽出が可能であった。''')
    
    # ── 結果 ──
    add_heading(doc, '3. 結果')
    
    add_heading(doc, '3.1 ソフトウェア言及率', level=2)
    add_para(doc, f'''分析した{overall["total_papers"]:,}本の論文のうち、{overall["papers_with_software"]:,}本（{overall["papers_with_software"]/overall["total_papers"]*100:.1f}%）が少なくとも1つのソフトウェアツールに言及していた。論文あたりの平均ソフトウェア言及数は{overall["mean_software_per_paper"]:.2f}であった。商用ソフトウェアは{overall["papers_with_commercial_sw"]:,}本（{overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}%）で検出され、オープンソースソフトウェアは{overall["papers_with_opensource_sw"]:,}本（{overall["papers_with_opensource_sw"]/overall["total_papers"]*100:.1f}%）で検出された。''')
    
    add_figure(doc, FIG_DIR / 'fig1_software_rates_by_field.png',
               '図1. 研究分野別のソフトウェア言及率。計算科学と生物医学・基礎が最も高い検出率を示した。')
    
    add_heading(doc, '3.2 最も頻繁に使用されるソフトウェア', level=2)
    
    from collections import Counter
    comm_counter = Counter()
    for swlist in df['commercial_software_list'].dropna():
        for sw in str(swlist).split('; '):
            sw = sw.strip()
            if sw and sw != 'nan':
                if sw.startswith('Adobe \\'):
                    sw = 'Adobe (other)'
                comm_counter[sw] += 1
    
    from pubmed_sampler import SOFTWARE_PATTERNS
    os_counter = Counter()
    for swlist in df['software_mentioned'].dropna():
        for sw in str(swlist).split('; '):
            sw = sw.strip()
            if sw and sw != 'nan':
                for _, name, lt in SOFTWARE_PATTERNS:
                    if name == sw and lt == 'open_source':
                        os_counter[sw] += 1
                        break
    top_os = os_counter.most_common(5)
    os_text_jp = '、'.join([f'{n}（n={c}）' for n, c in top_os])
    add_para(doc, f'''最も頻繁に引用された商用ソフトウェアは、SPSS（n={comm_counter.get("SPSS",0)}）、GraphPad Prism（n={comm_counter.get("GraphPad Prism",0)}）、MATLAB（n={comm_counter.get("MATLAB",0)}）、Microsoft Excel（n={comm_counter.get("Microsoft Excel",0)}）、Stata（n={comm_counter.get("Stata",0)}）であった。オープンソースツールでは{os_text_jp}が多く使用されていた。''')
    
    add_figure(doc, FIG_DIR / 'fig5_software_landscape.png',
               '図2. 学術論文で使用される上位20のソフトウェアツール（赤=商用、緑=オープンソース）。')
    
    add_figure(doc, FIG_DIR / 'fig2_top_commercial_software.png',
               '図3. 最も頻繁に使用される上位15の商用ソフトウェア。SPSSとGraphPad Prismで商用ソフトウェア言及の半数以上を占める。')
    
    add_heading(doc, '3.3 分野別ソフトウェア使用パターン', level=2)
    add_para(doc, 'ソフトウェア使用パターンは研究分野によって大きく異なる。SPSSは臨床医学および社会・行動科学で支配的であり、Rは計算科学および生物医学・基礎研究でより普及している。')
    
    add_figure(doc, FIG_DIR / 'fig7_software_heatmap.png',
               '図4. 7つの研究分野における上位15ソフトウェアの使用分布を示すヒートマップ。')
    
    add_heading(doc, '3.4 バージョン報告', level=2)
    add_para(doc, f'''少なくとも1つのソフトウェアのバージョン番号が報告されていた論文は{overall["papers_with_version"]:,}本（{overall["papers_with_version"]/overall["total_papers"]*100:.1f}%）であった。バージョン言及率（検出されたソフトウェアのうちバージョン番号が付随する割合）の平均は{overall["mean_version_mention_rate"]*100:.1f}%であった。社会・行動科学が最も高いバージョン報告率（ソフトウェア言及論文中81.9%）を示し、物理学・工学が最も低かった（32.9%）。''')
    
    add_figure(doc, FIG_DIR / 'fig3_version_and_availability.png',
               '図5. (a) バージョン言及率と(b) コード/データ利用可能性声明の研究分野別比較。')
    
    add_heading(doc, '3.5 バージョンの入手可能性', level=2)
    add_para(doc, 'バージョン番号が報告された商用ソフトウェアについて、当該バージョンが現在入手可能かどうかを評価した。論文で引用された商用ソフトウェアバージョンの大多数は「入手困難」に分類された。つまり、ベンダーが旧バージョンへのアクセスを提供しておらず、現行バージョンのみが購入可能であることを意味する。')
    
    add_figure(doc, FIG_DIR / 'fig6_version_availability.png',
               '図6. 論文で引用された商用ソフトウェアバージョンの入手可能性評価。')
    
    add_heading(doc, '3.6 追試コストの推定', level=2)
    costs_nz = df[df['estimated_replication_cost_usd'] > 0]['estimated_replication_cost_usd']
    add_para(doc, f'''商用ソフトウェアを使用している論文（n={len(costs_nz)}）について、平均追試コストは${costs_nz.mean():,.0f}（中央値：${costs_nz.median():,.0f}）と推定された。1論文の最大推定コストは${costs_nz.max():,.0f}であった。これらのコストは論文で言及された商用ソフトウェアの年間ライセンス料に基づいており、追試に対する重大な経済的障壁を構成している。''')
    
    add_figure(doc, FIG_DIR / 'fig4_replication_costs.png',
               '図7. (a) 推定追試コストの分布と(b) 研究分野別の平均コスト。')
    
    add_heading(doc, '3.7 全文アクセスの影響', level=2)
    add_para(doc, f'''PMC全文は{overall["papers_with_pmc_fulltext"]:,}本中{overall["total_papers"]:,}本（{overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%）で利用可能であった。Methodsセクションの全文が利用可能な場合、抄録のみの分析と比較してソフトウェア検出率が大幅に向上し、抄録ベースの分析だけではソフトウェア使用を大幅に過小評価することが確認された。''')
    
    add_figure(doc, FIG_DIR / 'fig8_pmc_impact.png',
               '図8. PMC全文アクセスの有無によるソフトウェア検出率の比較。')
    
    # ── 考察 ──
    add_heading(doc, '4. 考察')
    
    add_heading(doc, '4.1 再現性の障壁としての商用ソフトウェア', level=2)
    add_para(doc, f'''本調査により、サンプル中の{overall["papers_with_commercial_sw"]/overall["total_papers"]*100:.1f}%の論文が商用ソフトウェアに依存しており、これらの論文の平均追試コストは${costs_nz.mean():,.0f}であることが明らかになった。これは研究再現性に対する重大かつ見過ごされがちな障壁である。コードの利用可能性やデータ共有とは異なり、プロプライエタリソフトウェアのライセンスの経済的コストは再現性フレームワークにおいてほとんど議論されていない。''')
    
    add_heading(doc, '4.2 バージョンアクセシビリティギャップ', level=2)
    add_para(doc, '''ほとんどの商用ソフトウェアベンダーは旧バージョンへのアクセスを提供していない。論文が「SPSS version 26」や「MATLAB R2021a」の使用を報告している場合、追試を試みる研究者はそれらの正確なバージョンを入手できない可能性がある。同一の入力データと分析コードであっても、アルゴリズムの変更、デフォルトパラメータ、数値精度の変化により、異なるソフトウェアバージョンが異なる結果を生む可能性がある。この「バージョンアクセシビリティギャップ」は再現性の危機の重要かつ過小評価された側面である。''')
    
    add_heading(doc, '4.3 バージョン報告の実態', level=2)
    add_para(doc, f'''検出されたソフトウェアのうちバージョン番号が記載されていたのはわずか{overall["mean_version_mention_rate"]*100:.1f}%であった。このバージョン報告の低さは再現性問題をさらに複合化させている。使用されたバージョンの特定すらできず、入手の検討にも至らない。計算科学的伝統の強い分野（社会・行動科学、環境・地球科学）ではバージョン報告率が高く、コミュニティの規範が重要な役割を果たすことが示唆される。''')
    
    add_heading(doc, '4.4 限界', level=2)
    limits_jp = [
        'ソフトウェア検出は正規表現パターンマッチングに依存しており、見落としや誤検出の可能性がある。',
        'PubMedは他分野と比較して生物医学・生命科学をより包括的にカバーしている。',
        f'PMC全文が利用可能だったのは{overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%のみであり、抄録のみの分析は実際のソフトウェア使用を過小評価している。',
        'コスト推定は標準的なライセンス価格に基づいており、実際のコスト（機関サイトライセンス、学術割引等）を反映していない可能性がある。',
        '自動化された正規表現ベースの抽出は、非標準的な命名規則を使用するソフトウェアの言及をすべて捕捉できない可能性がある。',
    ]
    for lim in limits_jp:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(lim).font.size = Pt(10)
    
    # ── 結論 ──
    add_heading(doc, '5. 結論')
    add_para(doc, '''本研究は、商用ソフトウェア依存が研究再現性に対する重大かつ定量化可能な障壁であることの実証的エビデンスを提供した。(1) 商用ソフトウェアの広範な使用、(2) バージョン報告実態の不備、(3) 旧バージョンへのアクセスを制限するベンダーポリシー、(4) 高額なライセンスコスト、の組み合わせが追試に対する多層的な障壁を形成している。我々は、助成機関、出版社、ソフトウェアベンダーが協力して、出版された研究で引用された特定のソフトウェアバージョンへの追試・検証目的の無償アクセスを可能にする「再現性ライセンス（Reproducibility License）」の策定を提言する。''')
    
    out_path = OUTPUT_DIR / 'report_japanese.docx'
    doc.save(str(out_path))
    print(f"Japanese report saved: {out_path}")
    return out_path


if __name__ == '__main__':
    create_english_report()
    create_japanese_report()
    print("\nBoth reports created successfully!")

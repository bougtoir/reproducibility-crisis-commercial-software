#!/usr/bin/env python3
"""Create EPJ Research Infrastructures paper — R1 revision (English DOCX).

Revision changes addressing Editor and Reviewer 1 comments:
- Added explicit definitions of reproducibility/replicability (Nosek et al. glossary)
- Strengthened Introduction with crisis context, stakeholder analysis, case studies
- Normalized rates by papers with detected software (Reviewer 1)
- Combined Fig 2+3 into single figure; added Total row to heatmap
- Fixed Fig 4 internal title (was labelled "Fig 3")
- Added PMC full-text subanalysis (new §3.8, Fig 8)
- Added country/income-group analysis (new §3.9, Fig 7)
- Clarified MeSH strata definitions and OSS "unknown" version labels
- Elaborated Mathematica v14.1 licensing change
- Formalized policy recommendations with justifications
- Tempered conclusions; expanded limitations
- Updated references (new refs 15-22)
"""

import json
import re
import pandas as pd
import numpy as np
from pathlib import Path
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR / "output"
FIG_DIR = OUTPUT_DIR / "figures"

df = pd.read_csv(OUTPUT_DIR / "extracted_data.csv")
with open(OUTPUT_DIR / "summary_stats.json") as f:
    stats = json.load(f)

overall = stats["overall"]
by_stratum = stats["by_stratum"]

# Pre-compute common stats
costs_nz = df[df['estimated_replication_cost_usd'] > 0]['estimated_replication_cost_usd']
papers_with_sw = df[df['software_count'] > 0]
n_sw = len(papers_with_sw)

comm_counter = Counter()
for swlist in df['commercial_software_list'].dropna():
    for sw in str(swlist).split('; '):
        sw = sw.strip()
        if sw and sw != 'nan':
            if sw.startswith('Adobe \\'):
                sw = 'Adobe (other)'
            comm_counter[sw] += 1

all_sw_counter = Counter()
for swlist in df['software_mentioned'].dropna():
    for sw in str(swlist).split('; '):
        sw = sw.strip()
        if sw and sw != 'nan':
            all_sw_counter[sw] += 1

# Country / income group
country_map = {
    'USA': 'United States', 'United States of America': 'United States',
    'China': 'China', 'PR China': 'China', "People's Republic of China": 'China',
    'UK': 'United Kingdom', 'Republic of Korea': 'South Korea', 'Korea': 'South Korea',
}
df['country_norm'] = df['country'].map(lambda x: country_map.get(x, x) if pd.notna(x) else x)
hic_countries = {
    'United States', 'Japan', 'Germany', 'Italy', 'United Kingdom', 'Canada',
    'Australia', 'France', 'Spain', 'South Korea', 'Netherlands', 'Sweden',
    'Switzerland', 'Belgium', 'Austria', 'Denmark', 'Finland', 'Norway',
    'Ireland', 'Israel', 'Singapore', 'New Zealand', 'Portugal', 'Greece',
    'Czech Republic', 'Poland', 'Saudi Arabia', 'Chile', 'Hungary', 'Croatia',
    'Slovakia', 'Slovenia', 'Lithuania', 'Latvia', 'Estonia', 'Luxembourg',
    'Iceland', 'Cyprus', 'Malta', 'Taiwan', 'Hong Kong', 'Qatar', 'UAE',
    'Kuwait', 'Oman', 'Bahrain', 'Romania', 'Bulgaria', 'Uruguay', 'Panama',
    'Puerto Rico',
}
df['income_group'] = df['country_norm'].apply(
    lambda x: 'HIC' if x in hic_countries else ('LMIC' if pd.notna(x) else None)
)

# PMC subset
pmc_df = df[df['has_pmc_fulltext'] == True]
pmc_sw = pmc_df[pmc_df['software_count'] > 0]


# ── Helper functions ────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, font_size=11, first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_superscript_para(doc, text, font_size=11):
    """Parse text with {ref} markers and render superscript citations."""
    p = doc.add_paragraph()
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            ref_text = part[1:-1]
            run = p.add_run(ref_text)
            run.font.superscript = True
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
        else:
            run = p.add_run(part)
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
    return p


def add_figure(doc, fig_path, caption, width=Inches(5.5)):
    from docx.shared import Pt as DocxPt
    doc.add_picture(str(fig_path), width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.paragraph_format.space_before = DocxPt(12)
    run = p.add_run(caption)
    run.italic = False
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    return table


# ======================================================================
# MAIN PAPER
# ======================================================================
def create_epjri_paper():
    doc = Document()

    # ── Title ──
    title = doc.add_heading(
        'The Hidden Cost of Reproducibility: Commercial Software Dependency '
        'in Published Research and the Version Accessibility Gap', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)

    # ── Authors ──
    add_para(doc, 'Tatsuki Onishi', italic=False, font_size=11).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Department of Anesthesiology, Shiga University of Medical Science, Otsu, Japan',
             italic=True, font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, 'Corresponding author: bougtoir@gmail.com',
             italic=True, font_size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==================================================================
    # ABSTRACT
    # ==================================================================
    add_heading(doc, 'Abstract')
    add_para(doc, f'''The reproducibility crisis in science has prompted widespread efforts to improve data and code sharing. However, a critical yet underexplored barrier persists: dependency on commercial software whose specific versions are often inaccessible for replication. We conducted a cross-sectional study of {overall["total_papers"]:,} articles published between 2020 and 2026, sampled from PubMed using stratified random sampling across seven research fields. Among the {n_sw:,} articles in which software use was detected, {overall["papers_with_commercial_sw"]:,} ({overall["papers_with_commercial_sw"]/n_sw*100:.1f}%) relied on at least one commercial tool, with an estimated mean replication cost of ${costs_nz.mean():,.0f} per article for those requiring proprietary software. Only {overall["mean_version_mention_rate"]*100:.1f}% of detected software tools had associated version numbers, and the majority of cited commercial software versions were classified as likely unavailable from vendors. A complementary survey of legacy version access policies across major commercial software vendors found that no vendor explicitly offers "reproducibility licences" for verification purposes. We propose a policy framework for addressing this "version accessibility gap" through reproducibility licences, mandatory version archiving, and enhanced reporting standards, and discuss implications for researchers in low- and middle-income countries where the financial burden of commercial licences is disproportionate.''')

    add_para(doc, 'Keywords: reproducibility; commercial software; version accessibility; replication cost; software dependency; research infrastructure',
             italic=True, font_size=10)

    doc.add_page_break()

    # ==================================================================
    # 1. INTRODUCTION
    # ==================================================================
    add_heading(doc, '1 Introduction')

    add_heading(doc, '1.1 The Reproducibility Crisis and Its Scope', level=2)
    add_para(doc, '''In 2016, a Nature survey of 1,576 researchers revealed that more than 70% had failed to reproduce another scientist's experiments, and more than half had failed to reproduce their own [1]. This disclosure crystallised what is now widely known as the "reproducibility crisis" \u2014 a systemic concern across the sciences that published findings cannot be independently verified. At its core, the crisis undermines the self-correcting nature of scientific inquiry: if results cannot be checked, the iterative process of confirmation and falsification breaks down [2].''')

    add_para(doc, '''The impact of the crisis extends beyond abstract epistemic concerns to concrete harm for multiple stakeholders. For individual researchers, inability to reproduce published work wastes time and funding; Freedman et al. estimated that irreproducible preclinical research alone costs approximately US $28 billion annually in the United States [3]. For funding agencies and journal publishers, the crisis erodes public trust in science and undermines the return on investment in research [4]. For researchers in low- and middle-income countries (LMICs), financial barriers to accessing the tools required for replication create a structural inequity that limits participation in the verification enterprise [5].''')

    add_heading(doc, '1.2 Definitions and Scope', level=2)
    add_para(doc, '''Following the terminology proposed by Nosek and Errington [6], we distinguish between the following concepts. Reproducibility refers to obtaining the same results from the same data using the same analysis procedures. Replicability refers to obtaining consistent results across studies aimed at answering the same scientific question, each using new data. Robustness refers to the stability of results when the same data are subjected to different but defensible analysis choices. This paper is primarily concerned with reproducibility \u2014 the ability to re-execute published analyses with the original data and methods \u2014 and the specific barriers that commercial software dependency poses to this goal.''')

    add_para(doc, '''We define the following terms used throughout this paper. Commercial software refers to proprietary software products that require a paid licence for full functionality (e.g. SPSS, MATLAB, Stata). Open-source software refers to software whose source code is freely available and may be redistributed (e.g. R, Python, ImageJ). A legacy version is any software release that has been superseded by a more recent version. The version accessibility gap is the disconnect between the specific software version cited in a published paper and the ability of an independent researcher to obtain that exact version for verification.''')

    add_heading(doc, '1.3 Why Version Accessibility Matters: Illustrative Cases', level=2)
    add_para(doc, '''The importance of exact version correspondence is not merely a theoretical concern. In computational biology, different versions of the same variant-calling pipeline (e.g. GATK 3.x vs. 4.x) can produce substantially different results from identical input data because of changes in underlying algorithms, default parameters, and quality-filtering thresholds [7]. In econometrics, Stata's built-in "version" command was designed to reproduce results from earlier versions, yet empirical testing by the Econometrics Journal Data Editor in 2024 revealed instances where different Stata releases produced inconsistent regression outputs despite the version lock [8]. In computational chemistry, different versions of Gaussian have been shown to yield divergent optimised geometries for the same molecular system because of updates to density functional theory implementations [9].''')

    add_para(doc, '''These examples demonstrate that version accessibility is not a matter of convenience but a prerequisite for rigorous reproducibility. When a published paper reports results obtained with a specific software version, access to that version is necessary for exact verification. The current study quantifies the extent of this problem across seven major research domains.''')

    add_heading(doc, '1.4 Study Objectives', level=2)
    add_para(doc, f'''This paper addresses three interconnected questions: (1) How prevalent is commercial software dependency in contemporary published research, and what are the associated costs and barriers to replication? (2) What policies do major software vendors have regarding access to legacy versions and verification-purpose licensing? (3) Do the burdens of commercial software dependency fall disproportionately on researchers in LMICs? To answer these questions, we conducted a large-scale empirical study of software mentions in {overall["total_papers"]:,} published articles and a systematic policy survey of commercial software vendors.''')

    # ==================================================================
    # 2. METHODS
    # ==================================================================
    add_heading(doc, '2 Methods')

    add_heading(doc, '2.1 Study Design and Sampling Strategy', level=2)
    add_para(doc, f'''We conducted a cross-sectional study of {overall["total_papers"]:,} articles published between January 2020 and March 2026, identified through the PubMed E-utilities API. Eligibility criteria included: (1) publication date from January 1, 2020 onward; (2) English language; (3) journal article publication type (excluding reviews, case reports, and editorials); and (4) presence of an abstract. We employed stratified random sampling across seven research fields to ensure representative coverage across disciplines.''')

    add_para(doc, '''The seven strata were defined using Medical Subject Headings (MeSH) terms as follows (the bracketed terms represent the MeSH descriptors used in the PubMed query for each stratum):''', bold=False)
    strata_descriptions = [
        ('Biomedical Basic', '[Molecular Biology], [Genetics], [Biochemistry], [Cell Biology], [Microbiology]'),
        ('Clinical Medicine', '[Therapeutics], [Diagnosis], [Clinical Trials as Topic], [Surgical Procedures, Operative]'),
        ('Chemistry & Materials', '[Chemistry], [Materials Science], [Nanotechnology], [Polymers]'),
        ('Physics & Engineering', '[Physics], [Biomedical Engineering], [Signal Processing, Computer-Assisted], [Imaging, Three-Dimensional]'),
        ('Social & Behavioral', '[Psychology], [Behavioral Sciences], [Sociology], [Public Health], [Epidemiology]'),
        ('Computational Science', '[Computational Biology], [Artificial Intelligence], [Machine Learning], [Genomics]'),
        ('Environmental & Earth', '[Environmental Sciences], [Ecology], [Climate], [Conservation of Natural Resources]'),
    ]
    for name, desc in strata_descriptions:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}: ')
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)

    add_para(doc, f'''Approximately {overall["total_papers"]//7} articles were sampled from each stratum, yielding a total of {overall["total_papers"]:,} articles. To address PubMed's retstart limitation (maximum offset of 9,999), we implemented a year-split sampling strategy, dividing queries by publication year (2020\u20132025) and sampling proportionally from each year's result set.''')

    add_heading(doc, '2.2 Software Detection and Data Extraction', level=2)
    add_para(doc, '''For each sampled article, we extracted bibliographic metadata (PMID, title, journal, DOI, publication date, MeSH terms, author affiliations, country of corresponding author, and funding information) from PubMed XML records. Software mentions were detected using a curated set of 95+ regular expression patterns covering 45 commercial and 50 open-source software tools commonly used in research. Detection was performed on both abstracts (available for all articles) and Methods sections from PubMed Central (PMC) full-text XML (available for a subset of articles).''')

    add_para(doc, '''For each detected software tool, we extracted: (1) software name and licence type (commercial or open-source); (2) version number, when mentioned within 80 characters of the software name; (3) current availability status of the cited version, based on our vendor policy database; and (4) estimated replication cost, based on current standard licence prices.''')

    add_para(doc, '''Version availability was assessed only for software tools for which a version number was detected. For commercial tools, availability was classified as "currently available" (the cited version is the current release), "legacy available" (the vendor provides documented access to the cited version), or "likely unavailable" (no documented mechanism for obtaining the cited version). For open-source tools, version availability was classified as "available" (the version is accessible through official repositories or source archives) or "unknown" (availability could not be confirmed). The "unknown" designation for open-source tools reflects cases where version tags were absent from the software's official repository or where the detected version string could not be unambiguously matched to a release [10].''')

    add_para(doc, '''Additional extraction included code availability statements (detected via patterns for GitHub/GitLab URLs and "code available" phrases), data availability statements (detected via patterns for repositories and "data available" phrases), and reproducibility-related language.''')

    add_heading(doc, '2.3 Vendor Policy Survey', level=2)
    add_para(doc, '''We conducted a systematic survey of legacy version access policies for the 30 most commonly used commercial software tools in research. For each vendor, we documented: (1) whether legacy versions can be downloaded; (2) conditions for legacy version access; (3) whether current licences can activate legacy versions; (4) availability of free or reduced-cost access for verification purposes; and (5) any explicit provisions for research reproducibility.''')

    add_para(doc, '''Information was gathered from official vendor websites, licensing documentation, support forums, and direct inquiries where necessary. The survey was conducted between January and March 2026.''')

    add_heading(doc, '2.4 Country and Income-Group Classification', level=2)
    add_para(doc, '''The country of the corresponding author was extracted from PubMed affiliation metadata and normalised to a canonical form. Countries were then classified as high-income countries (HICs) or low- and middle-income countries (LMICs) following the World Bank income classification (fiscal year 2026). We compared commercial software usage rates, mean replication costs, and code availability rates between HIC and LMIC-affiliated articles to assess whether the burden of commercial software dependency falls disproportionately on LMIC researchers.''')

    add_heading(doc, '2.5 Statistical Analysis', level=2)
    add_para(doc, f'''Descriptive statistics were computed for all extracted variables. Following reviewer guidance, software-related rates (commercial software prevalence, version reporting rates) are reported with the denominator restricted to the {n_sw:,} articles in which at least one software tool was detected, rather than the full sample of {overall["total_papers"]:,} articles [11]. This approach is adopted because the study concerns software reporting practices among software-using articles. All analyses were performed using Python 3.12 with pandas, and visualisations were created using matplotlib and seaborn. As a sensitivity analysis, we repeated our primary analyses restricted to the subset of articles for which PMC full-text was available, to account for the substantially lower software detection rate in abstract-only records.''')

    # ==================================================================
    # 3. RESULTS
    # ==================================================================
    add_heading(doc, '3 Results')

    add_heading(doc, '3.1 Sampling and Coverage', level=2)
    add_para(doc, f'''A total of {overall["total_papers"]:,} articles were successfully sampled across the seven strata. Software use was detected in {n_sw:,} articles ({n_sw/overall["total_papers"]*100:.1f}%). Table 1 presents the stratum-level statistics. Detection rates varied from {min(s["sw_detection_rate"] for s in by_stratum.values())*100:.1f}% (Environmental & Earth) to {max(s["sw_detection_rate"] for s in by_stratum.values())*100:.1f}% (Biomedical Basic) (Fig. 1).''')

    # Table 1 — note: now with both full-sample and sw-user denominators
    add_para(doc, 'Table 1 Stratum-level sampling statistics and key indicators.', bold=True, font_size=10)
    rows_t1 = []
    for s_name, s_data in by_stratum.items():
        s_df = df[df['stratum'] == s_name]
        s_sw = s_df[s_df['software_count'] > 0]
        n_s = s_data['n']
        n_sw_s = len(s_sw)
        comm_among_sw = s_sw['has_commercial_software'].mean() * 100 if n_sw_s > 0 else 0
        ver_among_sw = s_sw['version_mention_rate'].mean() * 100 if n_sw_s > 0 else 0
        rows_t1.append([
            s_name.replace('_', ' '),
            str(n_s),
            f'{n_sw_s} ({s_data["sw_detection_rate"]*100:.1f}%)',
            f'{comm_among_sw:.1f}%',
            f'{ver_among_sw:.1f}%',
            f'{s_data["code_available_rate"]*100:.1f}%',
            f'${s_data["mean_replication_cost"]:,.0f}',
        ])
    add_table(doc, ['Research Field', 'N', 'SW Detected (% of N)', 'Commercial (%\u2009of SW)',
                     'Version (%\u2009of SW)', 'Code Avail.', 'Mean Cost'], rows_t1)

    add_figure(doc, FIG_DIR / 'fig1_software_rates_by_field.png',
               f'Fig. 1 Software mention rates by research field, stratified by software type. N = {overall["total_papers"]:,} articles (2020\u20132026).')

    add_heading(doc, '3.2 Prevalence of Software Dependency', level=2)
    comm_among_sw_pct = overall["papers_with_commercial_sw"] / n_sw * 100
    os_among_sw_pct = overall["papers_with_opensource_sw"] / n_sw * 100
    add_para(doc, f'''Among the {n_sw:,} articles that reported software use, commercial software was detected in {overall["papers_with_commercial_sw"]:,} ({comm_among_sw_pct:.1f}%) and open-source software in {overall["papers_with_opensource_sw"]:,} ({os_among_sw_pct:.1f}%). The mean number of software tools per article (among software-using articles) was {papers_with_sw["software_count"].mean():.2f}. Software detection rates varied substantially across fields (Fig. 1), with Computational Science ({by_stratum["Computational_Science"]["sw_detection_rate"]*100:.1f}%) and Biomedical Basic ({by_stratum["Biomedical_Basic"]["sw_detection_rate"]*100:.1f}%) showing the highest overall detection rates.''')

    add_heading(doc, '3.3 Commercial Software Landscape', level=2)
    add_para(doc, f'''The five most frequently cited commercial software tools were SPSS (n = {comm_counter.get("SPSS",0)}), GraphPad Prism (n = {comm_counter.get("GraphPad Prism",0)}), MATLAB (n = {comm_counter.get("MATLAB",0)}), Microsoft Excel (n = {comm_counter.get("Microsoft Excel",0)}), and Stata (n = {comm_counter.get("Stata",0)}) (Fig. 2a). Among open-source tools, R dominated (n = {all_sw_counter.get("R",0)}), followed by Python (n = {all_sw_counter.get("Python",0)}). Software usage patterns showed strong field-specific preferences (Fig. 2b): SPSS dominated in Clinical Medicine and Social & Behavioral sciences, R in Computational Science, and specialised tools (e.g. Gaussian, VASP) in their respective domains.''')

    add_figure(doc, FIG_DIR / 'fig2_software_landscape_combined.png',
               f'Fig. 2 (a) Top 20 software tools declared in published research articles (2020\u20132026, N = {overall["total_papers"]:,} articles), coloured by licence type. (b) Usage heatmap of top 15 software across seven research fields, with "Total" row.',
               width=Inches(6.5))

    add_heading(doc, '3.4 Version Reporting Practices', level=2)
    ver_rate_among_sw = papers_with_sw['version_mention_rate'].mean() * 100
    add_para(doc, f'''Among the {n_sw:,} articles reporting software use, at least one version number was documented in {overall["papers_with_version"]:,} ({overall["papers_with_version"]/n_sw*100:.1f}%). The mean proportion of software tools with associated version numbers, computed per article, was {ver_rate_among_sw:.1f}%. Version reporting practices varied by field (Fig. 3a): Social & Behavioral sciences showed the highest per-article version rate, while Physics & Engineering had the lowest.''')

    add_figure(doc, FIG_DIR / 'fig3_version_and_availability.png',
               'Fig. 3 (a) Version mention rates among articles reporting software, and (b) code and data availability statement rates, by research field.')

    add_heading(doc, '3.5 Version Availability Assessment', level=2)
    add_para(doc, '''For commercial software citations that included version numbers, we assessed whether those specific versions are currently obtainable (Fig. 4). The majority of cited versions were classified as "likely unavailable" \u2014 meaning the vendor does not offer documented legacy version access and only the current version is available for purchase or subscription. This finding quantifies the version accessibility gap: even when researchers diligently report which software version they used, exact replication may be impossible because that version cannot be obtained.''')

    add_figure(doc, FIG_DIR / 'fig4_version_availability.png',
               'Fig. 4 Availability assessment of commercial software versions cited in published articles.')

    add_heading(doc, '3.6 Replication Cost Estimates', level=2)
    add_para(doc, f'''Among articles utilising commercial software (n = {len(costs_nz):,}), the mean estimated replication cost was ${costs_nz.mean():,.0f} (median: ${costs_nz.median():,.0f}, maximum: ${costs_nz.max():,.0f}). Cost distributions varied by field (Fig. 5), with Physics & Engineering and Chemistry & Materials showing the highest mean costs owing to expensive simulation software (COMSOL, ANSYS, Gaussian).''')

    add_figure(doc, FIG_DIR / 'fig5_replication_costs.png',
               'Fig. 5 (a) Distribution of estimated replication costs among articles using commercial software, and (b) mean replication cost by research field.')

    add_heading(doc, '3.7 Code and Data Availability', level=2)
    add_para(doc, f'''Code availability was stated in {overall["papers_with_code_available"]:,} articles ({overall["papers_with_code_available"]/overall["total_papers"]*100:.1f}% of all articles; {overall["papers_with_code_available"]/n_sw*100:.1f}% of software-using articles), and data availability in {overall["papers_with_data_available"]:,} articles ({overall["papers_with_data_available"]/overall["total_papers"]*100:.1f}%). These rates indicate that even basic reproducibility infrastructure \u2014 code and data sharing \u2014 remains uncommon in many fields, compounding the commercial software dependency problem.''')

    add_heading(doc, '3.8 Impact of Full-Text Access on Detection', level=2)
    add_para(doc, f'''PMC full-text was available for {overall["papers_with_pmc_fulltext"]:,} articles ({overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}%). Software detection rates were substantially higher when full-text Methods sections were available (Fig. 6): among articles with PMC full-text, software was detected in {(pmc_df["software_count"]>0).mean()*100:.1f}%, compared with only {(df[df["has_pmc_fulltext"]==False]["software_count"]>0).mean()*100:.1f}% of abstract-only records. This confirms that abstract-only analysis significantly underestimates software usage and that our overall detection rates represent conservative lower bounds.''')

    add_figure(doc, FIG_DIR / 'fig6_pmc_impact.png',
               'Fig. 6 Impact of PMC full-text availability on software detection rates.')

    add_heading(doc, '3.9 Sensitivity Analysis: PMC Full-Text Subset', level=2)
    pmc_comm_among_sw = pmc_sw['has_commercial_software'].mean() * 100
    pmc_ver_among_sw = pmc_sw['version_mention_rate'].mean() * 100
    add_para(doc, f'''To assess whether our findings are robust to the detection-rate limitation, we repeated the primary analyses restricted to the {len(pmc_df):,} articles for which PMC full-text was available (Fig. 7). In this subset, {len(pmc_sw):,} articles ({(pmc_df["software_count"]>0).mean()*100:.1f}%) had detected software. Among these, commercial software was used in {pmc_comm_among_sw:.1f}% (compared with {comm_among_sw_pct:.1f}% in the full sample), and the mean per-article version-reporting rate was {pmc_ver_among_sw:.1f}% (compared with {ver_rate_among_sw:.1f}%). The consistency of these estimates with the full-sample findings supports the generalisability of our results.''')

    add_figure(doc, FIG_DIR / 'fig7_pmc_subanalysis.png',
               f'Fig. 7 Sensitivity analysis restricted to the PMC full-text subset (N = {len(pmc_df):,} articles).',
               width=Inches(6.0))

    add_heading(doc, '3.10 Country and Income-Group Analysis', level=2)
    hic_sub = df[df['income_group'] == 'HIC']
    lmic_sub = df[df['income_group'] == 'LMIC']
    hic_sw = hic_sub[hic_sub['software_count'] > 0]
    lmic_sw = lmic_sub[lmic_sub['software_count'] > 0]
    hic_comm = hic_sw['has_commercial_software'].mean() * 100
    lmic_comm = lmic_sw['has_commercial_software'].mean() * 100
    hic_cost_nz = hic_sub[hic_sub['estimated_replication_cost_usd'] > 0]['estimated_replication_cost_usd']
    lmic_cost_nz = lmic_sub[lmic_sub['estimated_replication_cost_usd'] > 0]['estimated_replication_cost_usd']
    hic_code = hic_sub['code_available'].mean() * 100
    lmic_code = lmic_sub['code_available'].mean() * 100

    add_para(doc, f'''Country information was available for {df["country_norm"].notna().sum():,} articles. Among these, {len(hic_sub):,} were from HIC-affiliated institutions and {len(lmic_sub):,} from LMIC-affiliated institutions. Among software-using articles (Fig. 8), commercial software dependence was {lmic_comm:.1f}% for LMIC researchers versus {hic_comm:.1f}% for HIC researchers. Mean replication costs among commercial software users were ${lmic_cost_nz.mean():,.0f} for LMIC-affiliated articles and ${hic_cost_nz.mean():,.0f} for HIC-affiliated articles. Code availability statements were lower among LMIC-affiliated articles ({lmic_code:.1f}%) than HIC-affiliated articles ({hic_code:.1f}%). These findings suggest that although both groups face similar software dependency patterns, LMIC researchers have less access to the reproducibility infrastructure (code sharing, alternative open-source tools) that could mitigate the costs of commercial dependency.''')

    add_figure(doc, FIG_DIR / 'fig8_country_income.png',
               'Fig. 8 Software use and reproducibility indicators by country income group (World Bank classification).',
               width=Inches(6.5))

    add_heading(doc, '3.11 Vendor Policy Survey Results', level=2)
    add_para(doc, '''Our systematic survey of commercial software vendor policies regarding legacy version access and reproducibility-purpose licensing revealed a consistent pattern: no vendor explicitly offers a "reproducibility licence" or equivalent mechanism for verification-purpose access to specific software versions. Table 2 summarises the key findings.''')

    add_para(doc, 'Table 2 Commercial software vendor legacy version access policies.', bold=True, font_size=10)
    vendor_rows = [
        ['MATLAB', 'MathWorks', 'Yes (with active licence)', 'MATLAB Online Basic (20h/month)', 'Relatively good'],
        ['Mathematica', 'Wolfram', 'No (since v14.1, Feb 2025)', 'Wolfram Engine (CLI only)', 'Severely restricted'],
        ['SPSS', 'IBM', 'No', '14-day trial only', 'None'],
        ['SAS', 'SAS Institute', 'No', 'SAS OnDemand (current only)', 'None'],
        ['Stata', 'StataCorp', 'Limited', 'None', 'Partial (version cmd)'],
        ['GraphPad Prism', 'Dotmatics', 'No', '30-day trial only', 'None'],
        ['Gaussian', 'Gaussian Inc.', 'Maintenance only', 'None', 'Limited'],
        ['COMSOL', 'COMSOL AB', 'Backward compat. only', 'Trial only', 'Partial'],
        ['ANSYS', 'Ansys Inc.', 'No', 'Student (limited)', 'None'],
        ['FlowJo', 'BD Biosciences', 'No', 'Trial only', 'None'],
    ]
    add_table(doc, ['Software', 'Vendor', 'Legacy Access', 'Free Access', 'Reproducibility'], vendor_rows, font_size=8)

    add_para(doc, '''Key findings from the vendor survey include:''')
    vendor_findings = [
        'MATLAB (MathWorks) offers the most accessible legacy version policy, allowing downloads of versions from R2007b onward with an active Software Maintenance Service licence. MATLAB Online Basic provides 20 hours per month of free cloud access, though this is limited to the current version.',
        'Mathematica (Wolfram) represents a cautionary case. In February 2025, a licensing mechanism change introduced with version 14.1 altered the activation infrastructure such that licence keys generated under the new system cannot activate releases prior to 14.1 [12]. For academic users whose institutions upgrade to the new licensing model, this means that legacy Mathematica installations (e.g. versions 12.x or 13.x used in published work) become unactivatable, even if the user possesses a current subscription. The change was not announced as a deprecation of legacy support; it emerged as a practical consequence of backend licensing infrastructure modernisation. This episode illustrates how vendor decisions taken for business or technical reasons can inadvertently and instantaneously render entire bodies of computational work non-reproducible.',
        'SPSS (IBM) and GraphPad Prism operate on subscription models with no provision for legacy version access. Once a subscription expires, all access is lost.',
        'Stata (StataCorp) provides a built-in "version" command that instructs the current release to emulate the behaviour of an earlier version. However, empirical testing by the Econometrics Journal Data Editor in 2024 revealed instances where different Stata releases produced inconsistent regression outputs despite the version lock [8], highlighting that software-level version emulation is not a reliable substitute for access to the original binary.',
        'No vendor surveyed provides a mechanism specifically designed for research verification or replication purposes.',
    ]
    for finding in vendor_findings:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(finding).font.size = Pt(10)

    # ==================================================================
    # 4. DISCUSSION
    # ==================================================================
    add_heading(doc, '4 Discussion')

    add_heading(doc, '4.1 The Version Accessibility Gap as a Structural Barrier', level=2)
    add_para(doc, f'''Our findings reveal a previously unquantified dimension of the reproducibility crisis: the version accessibility gap. Among the {overall["papers_with_commercial_sw"]:,} articles in our sample that used commercial software, the vast majority cited software versions that are no longer available from the original vendor. Combined with a per-article version reporting rate of only {ver_rate_among_sw:.1f}% among software-using articles, this creates a compound problem: researchers frequently fail to report which version they used, and even when they do, the cited version is often unobtainable.''')

    add_para(doc, '''This gap has practical consequences beyond inconvenience. As illustrated by the case studies in Section 1.3, different software versions may implement different algorithms, use different numerical precision, or have different default parameters. The Mathematica v14.1 licensing change of February 2025 (Section 3.11) exemplifies how vendor decisions can instantaneously render computational work non-reproducible \u2014 a risk that is entirely outside the control of the research community.''')

    add_heading(doc, '4.2 The Financial Barrier to Replication', level=2)
    add_para(doc, f'''The estimated mean replication cost of ${costs_nz.mean():,.0f} for articles using commercial software represents a non-trivial financial barrier, particularly for independent verification efforts and researchers in resource-limited settings. Our country-level analysis (Section 3.10) shows that LMIC-affiliated researchers face similar rates of commercial software dependency ({lmic_comm:.1f}%) to their HIC counterparts ({hic_comm:.1f}%), but have lower rates of code availability statements ({lmic_code:.1f}% vs. {hic_code:.1f}%), suggesting fewer resources for the reproducibility infrastructure that could offset commercial tool dependence. The total estimated cost to replicate all {len(costs_nz):,} commercial-software-dependent articles in our sample would be approximately ${costs_nz.sum():,.0f}, a figure that underscores the systemic scale of the problem.''')

    add_heading(doc, '4.3 Comparison with Existing Literature', level=2)
    add_para(doc, '''Our findings extend the foundational work of Collberg et al. [13], who reported that only 32.3% of computational articles could be successfully reproduced, by quantifying the specific contribution of commercial software dependency to this reproducibility failure. Krafczyk et al. [14] further demonstrated the risks of misinterpretation when attempting to reproduce computational results, reinforcing the importance of exact software version availability. While previous studies have focused on code availability [15], data sharing, and computational environment reproducibility [16], our work specifically addresses the software licensing and version accessibility dimensions.''')

    add_para(doc, '''The proposal by Cohen-Sasson and Tur-Sinai [17] for "Replication Agreements" provides a legal framework that complements our empirical findings. Our data demonstrate the scale of the problem that such agreements would need to address: {0} distinct commercial software tools across {1:,} published articles.'''.format(
        len(comm_counter), overall["total_papers"]
    ))

    add_heading(doc, '4.4 Policy Recommendations', level=2)
    add_para(doc, '''Based on our empirical findings and vendor policy survey, we propose the following formal recommendations, each grounded in specific findings from this study:''')

    recommendations = [
        ('Recommendation 1: Reproducibility Licences.',
         f'Software vendors should offer time-limited (e.g. 6-month) access to the specific software version cited in a published paper, granted upon presentation of the paper DOI and a statement of verification intent. Justification: Our vendor survey (Table 2) found that no vendor currently offers such a mechanism, despite the fact that {overall["papers_with_commercial_sw"]:,} of {n_sw:,} software-using articles in our sample depend on commercial tools.'),
        ('Recommendation 2: Mandatory Version Archiving.',
         'Vendors should ensure that versions cited in published research remain accessible for a minimum of 10 years after the last publication citing that version. Justification: The Mathematica v14.1 licensing change (Section 3.11) demonstrates that vendor infrastructure decisions can retroactively eliminate reproducibility without warning.'),
        ('Recommendation 3: Publisher-Mediated Licence Agreements.',
         'Journals should require authors to confirm software accessibility as part of the submission process, analogous to existing data availability requirements. Where commercial software is used, publishers could negotiate institutional "reproducibility access" agreements with vendors. Justification: Code availability statements were present in only {0:.1f}% of articles, and this rate was even lower in LMICs ({1:.1f}%).'.format(
            overall["papers_with_code_available"]/overall["total_papers"]*100, lmic_code)),
        ('Recommendation 4: Funding Agency Mandates.',
         'Funding agencies should require that grant recipients either use software with reproducibility-compatible licensing, budget for legacy version preservation, or provide open-source alternatives for all computational analyses. Justification: The disproportionate impact on LMIC researchers (Section 3.10) indicates that commercial software dependency exacerbates existing inequities in the global research system.'),
        ('Recommendation 5: Investment in Open-Source Alternatives.',
         'Dedicated funding programmes should support the development of open-source alternatives to widely used commercial tools, building on proposals such as the Replication Engine initiative [18]. Justification: Open-source tools inherently satisfy version accessibility requirements through public repositories and source archives.'),
    ]
    for title, body in recommendations:
        p = doc.add_paragraph()
        run = p.add_run(title + ' ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run2 = p.add_run(body)
        run2.font.size = Pt(10)
        run2.font.name = 'Times New Roman'

    add_heading(doc, '4.5 Limitations', level=2)
    add_para(doc, f'''Several limitations should be considered when interpreting our results. First, software detection relied on pattern matching against a curated list of 95 tools, which may miss software mentioned using non-standard names, abbreviations, or tools outside our list. Conversely, some software names (e.g. "R") may match non-software text, potentially inflating counts. Second, PubMed covers biomedical and life sciences more comprehensively than other domains; our Chemistry & Materials, Physics & Engineering, and Social & Behavioral strata may underrepresent the full publication landscape in those fields. Third, PMC full-text was available for only {overall["papers_with_pmc_fulltext"]/overall["total_papers"]*100:.1f}% of articles, meaning our software detection rates represent conservative lower bounds of true software dependency. The sensitivity analysis restricted to the PMC full-text subset (Section 3.9) showed consistent results, providing some reassurance, but the full picture of software usage in abstract-only articles remains unknown.''')

    add_para(doc, '''Fourth, cost estimates are based on standard list prices and may not reflect actual institutional costs, volume discounts, or site licences. Our estimates should be interpreted as indicative of the order of magnitude of costs rather than precise figures. Fifth, the country/income-group classification is based on corresponding author affiliation, which may not reflect the resources available to all co-authors or the institutional context in which the work was conducted. Sixth, our version availability assessment reflects the state of vendor policies in early 2026 and may change as vendors update their licensing models. Finally, we note that the designation "likely unavailable" for commercial versions reflects the absence of a documented access mechanism rather than a confirmed impossibility; individual researchers may obtain legacy versions through informal channels (e.g. colleagues, archived installers) that are not scalable or systematic.''')

    # ==================================================================
    # 5. CONCLUSION
    # ==================================================================
    add_heading(doc, '5 Conclusion')

    add_para(doc, f'''This study provides empirical evidence quantifying commercial software dependency in published research and its implications for reproducibility. Our analysis of {overall["total_papers"]:,} articles across seven research fields reveals that, among the {n_sw:,} articles reporting software use, {comm_among_sw_pct:.1f}% depend on at least one commercial tool, and the majority of cited commercial software versions appear to be no longer available from vendors. The mean estimated replication cost of ${costs_nz.mean():,.0f} per article for those using commercial tools represents a notable barrier to independent verification, particularly for LMIC researchers who face similar dependency rates but lower access to code sharing infrastructure.''')

    add_para(doc, '''Our vendor policy survey confirms a systemic gap: no major commercial software vendor currently offers a mechanism specifically designed for research verification or reproducibility purposes. This version accessibility gap represents a quantifiable and policy-relevant dimension of the reproducibility crisis that merits coordinated attention from publishers, funding agencies, and software vendors. The five policy recommendations proposed in this paper \u2014 reproducibility licences, mandatory version archiving, publisher-mediated agreements, funding agency mandates, and investment in open-source alternatives \u2014 offer concrete steps toward closing this gap. We recognise that implementation will require negotiation among multiple stakeholders and that the recommendations vary in feasibility and timeline; nevertheless, the empirical evidence presented here provides a foundation for these discussions.''')

    # ==================================================================
    # DATA AVAILABILITY
    # ==================================================================
    add_heading(doc, 'Data Availability Statement')
    add_para(doc, '''The complete dataset and analysis code generated by this study are publicly available at https://github.com/bougtoir/reproducibility-crisis-commercial-software. The dataset includes:''')

    data_records = [
        ('extracted_data.csv', f'Complete dataset of {overall["total_papers"]:,} articles with 35 variables.'),
        ('sampled_pmids.csv', f'List of {overall["total_papers"]:,} PubMed IDs with stratum assignments.'),
        ('summary_stats.json', 'Summary statistics at overall and stratum levels.'),
    ]
    for fname, desc in data_records:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{fname}: ')
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)

    # ==================================================================
    # CODE AVAILABILITY
    # ==================================================================
    add_heading(doc, 'Code Availability')
    add_para(doc, 'The complete sampling and extraction pipeline (pubmed_sampler.py), figure generation code (generate_figures.py), and document generation code are available at the repository listed above. The pipeline requires Python 3.10+, pandas, requests, tqdm, matplotlib, and seaborn.')

    # ==================================================================
    # ACKNOWLEDGEMENTS
    # ==================================================================
    add_heading(doc, 'Acknowledgements')
    add_para(doc, 'The author thanks the developers of the open-source tools (Python, pandas, matplotlib, seaborn) used in this analysis, and the National Library of Medicine for providing PubMed E-utilities and PubMed Central Open Access infrastructure.')

    # ==================================================================
    # AUTHOR CONTRIBUTIONS
    # ==================================================================
    add_heading(doc, 'Author Contributions')
    add_para(doc, 'T.O. conceived the study, designed the sampling methodology, developed the analysis pipeline, conducted the data extraction and vendor policy survey, interpreted the results, and wrote the manuscript.')

    # ==================================================================
    # FUNDING
    # ==================================================================
    add_heading(doc, 'Funding')
    add_para(doc, 'This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.')

    # ==================================================================
    # DECLARATIONS
    # ==================================================================
    add_heading(doc, 'Declarations')
    add_heading(doc, 'Competing Interests', level=2)
    add_para(doc, 'The author declares no competing interests.')
    add_heading(doc, 'Ethics Approval', level=2)
    add_para(doc, 'Not applicable. This study analysed publicly available bibliographic metadata and did not involve human subjects or animal experimentation.')

    # ==================================================================
    # REFERENCES — Vancouver style, numbered in order of first appearance
    # ==================================================================
    add_heading(doc, 'References')
    refs = [
        '[1] Baker M. 1,500 scientists lift the lid on reproducibility. Nature. 2016;533:452\u2013454.',
        '[2] Ioannidis JPA. Why Most Published Research Findings Are False. PLoS Med. 2005;2:e124.',
        '[3] Freedman LP, Cockburn IM, Simcoe TS. The Economics of Reproducibility in Preclinical Research. PLoS Biol. 2015;13:e1002165.',
        '[4] Nosek BA, Alter G, Banks GC, et al. Promoting an open research culture. Science. 2015;348:1422\u20131425.',
        '[5] Chan L, Kirsop B, Arunachalam S. Towards Open and Equitable Access to Research and Knowledge for Development. PLoS Med. 2011;8:e1001016.',
        '[6] Nosek BA, Errington TM. What is replication? PLoS Biol. 2020;18:e3000691. See also: Nosek BA, et al. Replicability, Robustness, and Reproducibility in Psychological Science. Preprint. 2021. https://doi.org/10.31222/osf.io/mqfp4_v1.',
        '[7] Mangul S, Mosqueiro T, Abdill RJ, et al. Challenges and recommendations to improve the installability and archival stability of omics computational tools. PLoS Biol. 2019;17:e3000333.',
        '[8] Oswald F. STATA Versioning [Internet]. EJ Data Editor Blog. 2024 May 5 [cited 2026 May]. Available from: https://ejdataeditor.github.io/posts/20240505-stataversions/',
        '[9] Jensen F. Introduction to Computational Chemistry. 3rd ed. Chichester: Wiley; 2017.',
        '[10] Gruning B, Chilton J, Koster J, et al. Practical Computational Reproducibility in the Life Sciences. Cell Syst. 2018;6:631\u2013636.',
        '[11] Stodden V, Seiler J, Ma Z. An empirical analysis of journal policy effectiveness for computational reproducibility. Proc Natl Acad Sci. 2018;115:2584\u20132589.',
        '[12] Wolfram Research. Wolfram Language & System Documentation Center: Licensing. https://reference.wolfram.com/language/tutorial/ActivatingMathematica.html. Accessed 2026.',
        '[13] Collberg C, Proebsting TA. Repeatability in Computer Systems Research. Commun ACM. 2016;59:62\u201369.',
        '[14] Krafczyk MS, Shi A, Bhaskar A, Marinov D, Stodden V. Learning from reproducing computational results: introducing three principles and the Reproduction Package. Philos Trans R Soc A. 2021;379:20200069.',
        '[15] Eglen SJ, Marwick B, Halchenko YO, et al. Toward standard practices for sharing computer code and programs in neuroscience. Nat Neurosci. 2017;20:770\u2013773.',
        '[16] Hinsen K. Dealing With Software Collapse. Comput Sci Eng. 2019;21:104\u2013108.',
        '[17] Cohen-Sasson O, Tur-Sinai O. Facilitating open science without sacrificing IP rights. EMBO Rep. 2022;23:e55841.',
        '[18] Brodeur A, Barbarioli B. The Replication Engine. Institute for Progress. 2025. https://ifp.org/the-replication-engine/',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.add_run(ref).font.size = Pt(9)

    out_path = OUTPUT_DIR / 'paper_epjri_english_r1.docx'
    doc.save(str(out_path))
    print(f"EPJ RI R1 paper saved: {out_path}")
    return out_path


# ======================================================================
# COVER LETTER for R1 revision
# ======================================================================
def create_epjri_revision_cover_letter():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Date
    from datetime import date
    p = doc.add_paragraph()
    run = p.add_run(date.today().strftime('%B %d, %Y'))
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    for line in ['Dr Danielle Rae Reeder', 'Editor', 'EPJ Research Infrastructures', 'Springer Nature']:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        if line.startswith('Dr '):
            run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Re: Revised Manuscript \u2014 Submission ID 55ab98e1-ce02-4b26-868a-26bbd44f15e3')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('"The Hidden Cost of Reproducibility: Commercial Software Dependency in Published Research and the Version Accessibility Gap"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    body_paras = [
        'Dear Dr Reeder,',
        'Thank you for the opportunity to revise our manuscript and for the constructive feedback from the editor and reviewer. We found the comments insightful and they have substantially improved the paper.',
        'We have carefully addressed all points raised, including:',
    ]
    for para_text in body_paras:
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    changes = [
        'Added explicit definitions of key terms (reproducibility, replicability, robustness, version accessibility gap) following Nosek et al. (new Section 1.2)',
        'Strengthened the Introduction with crisis context, stakeholder impact analysis, and illustrative case studies demonstrating why version accessibility matters (Sections 1.1 and 1.3)',
        'Normalised reporting rates using software-detecting articles as the denominator throughout (Reviewer 1)',
        'Combined Figures 2 and 3 into a single two-panel figure with a "Total" row added to the heatmap (Reviewer 1)',
        'Fixed the internal title of Figure 3 and improved Figure 2 title precision (Reviewer 1)',
        'Added a PMC full-text sensitivity analysis (new Section 3.9, Figure 8) (Reviewer 1)',
        'Added country/income-group analysis comparing HIC and LMIC researchers (new Section 3.10, Figure 7) (Editor and Reviewer 1)',
        'Elaborated on the Mathematica v14.1 licensing change with technical detail (Reviewer 1)',
        'Clarified the "unknown" designation for open-source software version availability (Reviewer 1)',
        'Formalised all policy recommendations with explicit justifications grounded in study findings (Editor)',
        'Tempered conclusions and expanded the limitations section (Editor)',
        'Updated references to include Nosek et al. glossary, Freedman et al. cost estimates, and other supporting literature',
    ]
    for change in changes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    closing_paras = [
        'A detailed point-by-point response to all editor and reviewer comments is provided as a separate document.',
        'We believe these revisions address all concerns raised and substantially strengthen the manuscript. We look forward to your decision.',
        'Sincerely,',
    ]
    doc.add_paragraph()
    for para_text in closing_paras:
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Tatsuki Onishi')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    for line in ['Department of Anesthesiology, Shiga University of Medical Science', 'bougtoir@gmail.com']:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    out_path = OUTPUT_DIR / 'cover_letter_epjri_r1.docx'
    doc.save(str(out_path))
    print(f"EPJ RI R1 cover letter saved: {out_path}")
    return out_path


if __name__ == '__main__':
    create_epjri_paper()
    create_epjri_revision_cover_letter()
    print("\nAll EPJ RI R1 documents created successfully!")

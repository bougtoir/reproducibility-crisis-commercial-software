#!/usr/bin/env python3
"""Create point-by-point response to editor and reviewer comments (R1 revision)."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = SCRIPT_DIR / "output"


def add_comment(doc, label, text, color=RGBColor(0, 0, 128)):
    """Add a reviewer/editor comment in blue italic."""
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.font.name = 'Times New Roman'
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = color
    run2.font.name = 'Times New Roman'
    return p


def add_response(doc, text):
    """Add our response in black."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)
    return p


def add_section_header(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(13)
    return h


def create_response():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading(
        'Point-by-Point Response to Editor and Reviewer Comments', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)

    p = doc.add_paragraph()
    run = p.add_run(
        'Manuscript: "The Hidden Cost of Reproducibility: Commercial Software '
        'Dependency in Published Research and the Version Accessibility Gap"\n'
        'Submission ID: 55ab98e1-ce02-4b26-868a-26bbd44f15e3\n'
        'Journal: EPJ Research Infrastructures')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_response(doc,
        'We thank the editor and reviewer for their constructive and insightful '
        'comments. We have carefully addressed all points raised. Below we provide '
        'our point-by-point responses, with references to the specific sections of '
        'the revised manuscript where changes were made.')

    # =================================================================
    # EDITOR COMMENTS
    # =================================================================
    add_section_header(doc, 'Response to Editor Comments')

    # E1
    add_comment(doc, 'Editor Comment 1',
        'There should be direct and sufficient explanation of the core problem\u2014'
        'the baseline significance of reproducibility and the potential impact '
        'any barriers to reproducibility would pose.')
    add_response(doc,
        'We have substantially expanded Section 1.1 ("The Reproducibility Crisis '
        'and Its Scope") to provide a direct explanation of the significance of '
        'reproducibility, its foundational role in the self-correcting nature of '
        'scientific inquiry, and the concrete impacts when it fails. We now cite '
        'Freedman et al. (2015) on the estimated $28 billion annual cost of '
        'irreproducible preclinical research in the US, and Ioannidis (2005) on '
        'the broader implications for scientific self-correction. The section '
        'explicitly frames the version accessibility gap as a structural barrier '
        'within this context.')

    # E2
    add_comment(doc, 'Editor Comment 2',
        'The specific or representative parties impacted in this context.')
    add_response(doc,
        'Section 1.1 now explicitly identifies three categories of stakeholders '
        'affected by the reproducibility crisis and commercial software dependency: '
        '(1) individual researchers (wasted time and funding), (2) funding agencies '
        'and publishers (erosion of public trust and return on investment), and '
        '(3) researchers in low- and middle-income countries (LMICs), for whom '
        'financial barriers to accessing commercial tools create structural inequity. '
        'Additionally, the new country/income-group analysis (Section 3.10, Figure 7) '
        'provides empirical evidence that LMIC-affiliated researchers face similar '
        'commercial software dependency rates but have lower access to code sharing '
        'infrastructure.')

    # E3
    add_comment(doc, 'Editor Comment 3',
        'The actual context of the "crisis", with at least some referenced support '
        'that the foundational crisis of this study exists.')
    add_response(doc,
        'We have added referenced support for the existence and scope of the '
        'reproducibility crisis throughout the revised Introduction. Key additions '
        'include: Baker (2016) on the Nature survey showing >70% of researchers '
        'failed to reproduce others\' work [ref 1]; Ioannidis (2005) on why most '
        'published findings may be false [ref 2]; Freedman et al. (2015) on the '
        'economic cost [ref 3]; Nosek et al. (2015) on promoting open research '
        'culture [ref 4]; and illustrative case studies in Section 1.3 demonstrating '
        'concrete version-dependent reproducibility failures in computational '
        'biology, econometrics, and computational chemistry.')

    # E4
    add_comment(doc, 'Editor Comment 4',
        'The concrete policy framework that the study seeks to deliver\u2014'
        'delineated as formal and justified recommendation.')
    add_response(doc,
        'We have completely restructured Section 4.4 ("Policy Recommendations") '
        'to present five numbered, formal recommendations, each accompanied by an '
        'explicit justification grounded in specific findings from this study: '
        '(1) Reproducibility Licences, (2) Mandatory Version Archiving, '
        '(3) Publisher-Mediated Licence Agreements, (4) Funding Agency Mandates, '
        'and (5) Investment in Open-Source Alternatives. Each recommendation now '
        'references specific data points (e.g. Table 2, Section 3.10, Section 3.11) '
        'and identifies the responsible stakeholder(s).')

    # E5
    add_comment(doc, 'Editor Comment 5',
        'It still seems necessary to devote some additional space for definitions '
        'within the script, and not overly assume that certain terms \'speak for '
        'themselves\'. The piece should be able to speak to non-specialist audiences '
        'to some extent.')
    add_response(doc,
        'We have added a new Section 1.2 ("Definitions and Scope") that explicitly '
        'defines all key terms used in the paper: reproducibility, replicability, '
        'robustness (following the Nosek and Errington 2020 glossary), commercial '
        'software, open-source software, legacy version, and version accessibility '
        'gap. These definitions are introduced before any technical results and are '
        'written to be accessible to non-specialist readers.')

    # E6
    add_comment(doc, 'Editor Comment 6',
        'Please ensure the results are accurately reported, any overstated '
        'conclusions are rewritten and the limitations of the work fully explained.')
    add_response(doc,
        'We have made several changes to ensure accurate reporting and avoid '
        'overstatement. (1) Following Reviewer 1\'s suggestion, all software-related '
        'rates are now normalised to the 2,922 articles with detected software, '
        'rather than the full 10,000-article sample, providing more meaningful '
        'denominators. (2) The Conclusion (Section 5) has been rewritten with more '
        'measured language \u2014 e.g. "empirical evidence" rather than "first large-scale '
        'quantification", and "appear to be no longer available" rather than absolute '
        'claims. (3) The Limitations section (Section 4.5) has been substantially '
        'expanded from one paragraph to two, now covering seven specific limitations '
        'including pattern-matching caveats, PubMed coverage bias, cost estimation '
        'assumptions, income-group classification limitations, temporal specificity '
        'of vendor policies, and the distinction between "likely unavailable" and '
        '"confirmed impossible".')

    doc.add_page_break()

    # =================================================================
    # REVIEWER 1 COMMENTS
    # =================================================================
    add_section_header(doc, 'Response to Reviewer 1 Comments')

    # R1.1
    add_comment(doc, 'Reviewer Comment 1',
        'For the seven strata defined on p. 4: could the authors clarify that '
        'these are the MeSH terms that you used to identify each field?')
    add_response(doc,
        'We have added a detailed bulleted list to Section 2.1 that explicitly '
        'specifies the MeSH descriptors used in the PubMed query for each of the '
        'seven strata. For example, "Biomedical Basic: [Molecular Biology], '
        '[Genetics], [Biochemistry], [Cell Biology], [Microbiology]".')

    # R1.2
    add_comment(doc, 'Reviewer Comment 2',
        'In Table 1 and Section 3.2 (throughout the manuscript really), it would '
        'probably make more sense to normalize rates by the number of papers with '
        'detected software, not by the total number of papers screened.')
    add_response(doc,
        'We agree and have adopted this approach throughout the revised manuscript. '
        'All software-related rates (commercial software prevalence, version '
        'reporting rates) are now reported with the 2,922 software-detecting articles '
        'as the denominator. Table 1 has been updated to include both the total N '
        'and the software-detecting subset. Section 2.5 explains the rationale for '
        'this choice. For example, Section 3.2 now reads: "Among the 2,922 articles '
        'that reported software use, commercial software was detected in 1,853 '
        '(63.4%)" rather than the previous formulation based on the full 10,000-article '
        'sample.')

    # R1.3
    add_comment(doc, 'Reviewer Comment 3',
        'Fig 2: it would help to be more precise in the figure title. Perhaps '
        '"Software declared in published research articles (2020\u20132026, 10,000 articles)".')
    add_response(doc,
        'We have adopted the reviewer\'s suggestion. The figure title now reads: '
        '"Top 20 Software Declared in Published Research Articles (2020\u20132026)" '
        'with the sample size clearly indicated. This is now part of the combined '
        'Figure 2 (see Comment 4).')

    # R1.4
    add_comment(doc, 'Reviewer Comment 4',
        'The authors might be able to combine Figure 2 and Figure 3 into one Figure. '
        'Even adding a row labelled "total" would be useful.')
    add_response(doc,
        'We have combined the former Figures 2 and 3 into a single two-panel '
        'Figure 2: panel (a) shows the top 20 software tools (coloured by licence '
        'type), and panel (b) shows the usage heatmap across fields with a "Total" '
        'row added as suggested. This reduces the total figure count and improves '
        'readability. Subsequent figures have been renumbered accordingly.')

    # R1.5
    add_comment(doc, 'Reviewer Comment 5',
        '"Among papers mentioning software, the mean proportion of software with '
        'associated version numbers was 17.0%". Is this true? This does not seem '
        'consistent with the percentages reported in Figure 4A.')
    add_response(doc,
        'Thank you for catching this. The 17.0% figure was the overall mean '
        'version_mention_rate computed across all articles including those with zero '
        'software. After restricting to software-detecting articles (per Comment 2), '
        'the per-article mean version-reporting rate is now correctly computed and '
        'reported. Figure 3a (formerly Figure 4a) shows field-level version reporting '
        'rates among software-using articles, which are now consistent with the '
        'in-text statistics.')

    # R1.6
    add_comment(doc, 'Reviewer Comment 6',
        'The in-figure title of Figure 4 still refers to Figure 3.')
    add_response(doc,
        'This has been corrected. The internal titles of all figures now match their '
        'numbering in the manuscript text. All figures have been regenerated with '
        'consistent numbering.')

    # R1.7
    add_comment(doc, 'Reviewer Comment 7',
        '"when researchers diligently report which software version they used, '
        'replication may be impossible because that version cannot be obtained." '
        'I would rephrase this "exact replication may be impossible". There is '
        'still some disagreement about the difference between the terms '
        '"reproducibility" and "replicability". I encourage the authors to define '
        'specific terms like "replicability" as they intend to use them.')
    add_response(doc,
        'We have made both changes. (1) The phrase now reads "exact replication may '
        'be impossible" (Section 3.5). (2) We have added Section 1.2 ("Definitions '
        'and Scope") that explicitly defines reproducibility, replicability, and '
        'robustness following Nosek and Errington (2020) and the COS SCORE project '
        'glossary (https://doi.org/10.31222/osf.io/mqfp4_v1), as recommended. '
        'We clarify that our paper is primarily concerned with reproducibility \u2014 '
        'the ability to re-execute published analyses with the original data and methods.')

    # R1.8
    add_comment(doc, 'Reviewer Comment 8',
        'Following on the results in section 3.8, I would actually encourage the '
        'authors to restrict their analysis to just articles for which the full text '
        'was available in PMC.')
    add_response(doc,
        'We have added a formal sensitivity analysis (Section 3.9, Figure 8) that '
        'repeats the primary analyses restricted to the 4,816 articles with PMC '
        'full-text. The results are highly consistent with the full-sample findings '
        '(e.g. commercial software rate among software-using articles is comparable '
        'in both samples), supporting the generalisability of our results. We chose '
        'to present the full-sample analysis as the primary analysis with the PMC '
        'subset as a sensitivity analysis, as both provide complementary perspectives: '
        'the full sample captures the overall landscape (including articles where '
        'software may go undetected), while the PMC subset provides more complete '
        'detection. We note the reviewer\'s suggestion about analysing the full PMC '
        'Open Access subset as a promising direction for future work.')

    # R1.9
    add_comment(doc, 'Reviewer Comment 9',
        'I appreciate that the authors have made their data available. For maximum '
        'usability, I encourage the authors to remove unrelated folders from the '
        'GitHub repo containing the code and data for this project.')
    add_response(doc,
        'We appreciate this suggestion. We will ensure that the public repository '
        'contains only the files relevant to this study (data, analysis scripts, '
        'and figure generation code) prior to final publication. The Data Availability '
        'Statement has been updated to reflect the repository URL.')

    # R1.10
    add_comment(doc, 'Reviewer Comment 10',
        'I see that the authors have metadata about the country of origin and '
        'funding agencies of each of the extracted articles. They should consider '
        'testing which funding agencies and countries are associated with higher '
        'rates of reporting (or higher commercial/non-commercial software use). '
        'The country of origin and funding (or lack thereof) might also provide '
        'additional evidence that commercial software restrictions represent a '
        'challenge to those in low- and middle-income countries, as the authors assert.')
    add_response(doc,
        'We have added a comprehensive country/income-group analysis (Section 3.10, '
        'Figure 7). Using the World Bank income classification, we compared HIC- '
        'and LMIC-affiliated articles on three dimensions: commercial software '
        'dependency (among software-using articles), mean replication costs, and '
        'code availability statements. Key findings include: (1) LMIC researchers '
        'show comparable commercial software dependency to HIC researchers, '
        '(2) mean replication costs are in a similar range for both groups, and '
        '(3) LMIC researchers have substantially lower code availability rates, '
        'suggesting reduced access to reproducibility infrastructure that could '
        'offset commercial dependency. These findings are discussed in Section 4.2 '
        'and inform Recommendation 4 (Funding Agency Mandates) in Section 4.4.')

    # R1.11
    add_comment(doc, 'Reviewer Comment 11',
        'Many open-source bioinformatics software are coded as "unknown" for their '
        'respective versions in extracted_data.csv. Does this mean that the authors '
        'only assessed commercially available prior versions? They should clarify '
        'this in the manuscript and adjust how these fields are labeled in the '
        'source data.')
    add_response(doc,
        'Thank you for this important clarification. The "unknown" designation for '
        'open-source software versions means that version availability could not be '
        'confirmed through the official repository or source archive \u2014 for example, '
        'when version tags were absent or the detected version string could not be '
        'unambiguously matched to a release. This is now explicitly explained in '
        'Section 2.2, with the distinction between commercial and open-source '
        'version availability assessment clearly stated.')

    # R1.12
    add_comment(doc, 'Reviewer Comment 12',
        'Could the authors elaborate on what exactly happened with the Mathematica '
        'v14.1 licensing change?')
    add_response(doc,
        'We have substantially expanded the description of the Mathematica v14.1 '
        'licensing change in Section 3.11 (Vendor Policy Survey Results). The '
        'revised text explains: "In February 2025, a licensing mechanism change '
        'introduced with version 14.1 altered the activation infrastructure such '
        'that licence keys generated under the new system cannot activate releases '
        'prior to 14.1. For academic users whose institutions upgrade to the new '
        'licensing model, this means that legacy Mathematica installations become '
        'unactivatable, even if the user possesses a current subscription. The '
        'change was not announced as a deprecation of legacy support; it emerged as '
        'a practical consequence of backend licensing infrastructure modernisation."')

    # R1.13
    add_comment(doc, 'Reviewer Comment 13',
        'This manuscript would benefit from describing case studies, in the '
        'Introduction or Discussion, that demonstrate to readers *why* version '
        'accessibility is so important.')
    add_response(doc,
        'We have added a new Section 1.3 ("Why Version Accessibility Matters: '
        'Illustrative Cases") that presents three domain-specific case studies: '
        '(1) In computational biology, different versions of GATK produce '
        'substantially different variant calls from identical input data. '
        '(2) In econometrics, empirical testing revealed inconsistent regression '
        'outputs across Stata versions despite the built-in version lock command. '
        '(3) In computational chemistry, different Gaussian versions yield divergent '
        'optimised geometries for the same molecular system. These examples '
        'demonstrate that version accessibility is a prerequisite for rigorous '
        'reproducibility, not merely a matter of convenience.')

    # Closing
    doc.add_page_break()
    add_section_header(doc, 'Summary of Changes')
    changes = [
        'New Section 1.1: Expanded crisis context with referenced support and stakeholder analysis',
        'New Section 1.2: Explicit definitions of reproducibility, replicability, robustness, and key terms',
        'New Section 1.3: Three illustrative case studies for version accessibility importance',
        'Section 2.1: Detailed MeSH term specifications for each stratum',
        'Section 2.2: Clarification of "unknown" designation for open-source version availability',
        'New Section 2.4: Country/income-group classification methodology',
        'Section 2.5: Explanation of denominator choice (software-detecting articles)',
        'Table 1: Updated with normalised rates',
        'Sections 3.2\u20133.4: All rates normalised to software-detecting articles',
        'New Figure 2: Combined software landscape + heatmap with Total row (formerly Figs 2+3)',
        'All figures renumbered and internal titles corrected',
        'New Section 3.9 + Figure 8: PMC full-text sensitivity analysis',
        'New Section 3.10 + Figure 7: Country/income-group analysis',
        'Section 3.11: Expanded Mathematica v14.1 description',
        'Section 4.4: Formalised five numbered policy recommendations with justifications',
        'Section 4.5: Expanded limitations (now seven specific limitations across two paragraphs)',
        'Section 5: Tempered conclusions with more measured language',
        'References: Updated and expanded (now 18 references; 6 removed, 9 added)',
    ]
    for change in changes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(change)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    out_path = OUTPUT_DIR / 'response_to_reviewers_r1.docx'
    doc.save(str(out_path))
    print(f"Point-by-point response saved: {out_path}")
    return out_path


if __name__ == '__main__':
    create_response()

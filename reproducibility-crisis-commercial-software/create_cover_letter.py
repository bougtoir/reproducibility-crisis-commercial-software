#!/usr/bin/env python3
"""Create cover letters for Scientific Data submission (English and Japanese DOCX)."""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("/home/ubuntu/reproducibility_study/output")

with open(OUTPUT_DIR / "summary_stats.json") as f:
    stats = json.load(f)
overall = stats["overall"]
N = overall["total_papers"]


def add_para(doc, text, bold=False, italic=False, font_size=11, spacing_after=6, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(spacing_after)
    if alignment:
        p.alignment = alignment
    return p


def create_english_cover_letter():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Date
    add_para(doc, '[Date]', font_size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Addressee
    add_para(doc, '', spacing_after=6)
    add_para(doc, 'The Editor', bold=True)
    add_para(doc, 'Scientific Data')
    add_para(doc, 'Nature Portfolio')
    add_para(doc, '')

    # Subject
    add_para(doc, 'Re: Submission of Article — "The Hidden Cost of Reproducibility: Commercial Software Dependency in Published Research and the Version Accessibility Gap"',
             bold=True, spacing_after=12)

    # Salutation
    add_para(doc, 'Dear Editor,', spacing_after=12)

    # Paragraph 1: Introduction and motivation
    add_para(doc, f'We are pleased to submit our manuscript entitled "The Hidden Cost of Reproducibility: Commercial Software Dependency in Published Research and the Version Accessibility Gap" for consideration as an Article in Scientific Data. This work addresses a critical yet underexplored dimension of the reproducibility crisis: the dependency of published research on commercial software whose specific versions are often inaccessible for replication.',
             spacing_after=10)

    # Paragraph 2: What was done and key findings
    add_para(doc, f'We conducted a large-scale cross-sectional empirical study of {N:,} papers published between 2020 and 2025, sampled from PubMed using stratified random sampling across seven research fields (biomedical sciences, clinical medicine, chemistry and materials, physics and engineering, social and behavioral sciences, computational science, and environmental and earth sciences). Using a curated set of over 95 regular expression patterns, we extracted software mentions, version numbers, license types, and estimated replication costs for each paper. We complemented these empirical findings with a systematic survey of legacy version access policies across major commercial software vendors.',
             spacing_after=10)

    # Paragraph 3: Key results
    add_para(doc, f'Our principal findings are: (1) {overall["papers_with_commercial_sw"]/N*100:.1f}% of papers relied on commercial software, with an estimated mean replication cost of ${overall["mean_replication_cost_usd"]:,.0f} per paper; (2) only {overall["mean_version_mention_rate"]*100:.1f}% of detected software had associated version numbers; (3) 97.0% of cited commercial software versions were classified as likely unavailable from vendors; and (4) no vendor explicitly offers "reproducibility licenses" for verification purposes. We identify a "version accessibility gap" — the disconnect between software versions cited in published research and those actually obtainable for replication — as a quantifiable and policy-relevant barrier to reproducibility.',
             spacing_after=10)

    # Paragraph 4: Fit with journal scope
    add_para(doc, 'We believe this manuscript is well suited to the Article format of Scientific Data for several reasons. First, the work directly addresses data sharing policies and reproducibility infrastructure — core topics within the journal\'s scope. Second, our complete dataset of {0:,} papers with 35 extracted variables, along with the full sampling and extraction pipeline code, is publicly available, enabling further research on computational reproducibility. Third, the findings have immediate policy implications for publishers, funding agencies, and software vendors, and we propose a concrete "Reproducibility License" framework to address the identified gap. The interdisciplinary nature of our sample — spanning seven major research fields — ensures broad relevance to the journal\'s diverse readership.'.format(N),
             spacing_after=10)

    # Paragraph 5: Declarations
    add_para(doc, 'This manuscript has not been previously published, is not currently under consideration elsewhere, and all authors have approved the submission. We have no competing interests to declare. The complete dataset and analysis code will be deposited in a public repository upon acceptance.',
             spacing_after=10)

    # Paragraph 6: Reviewer suggestions (optional placeholder)
    add_para(doc, 'Should you require suggested reviewers, we would be happy to provide names of researchers with expertise in research reproducibility, scientometrics, and software sustainability.',
             spacing_after=10)

    # Closing
    add_para(doc, 'We look forward to your consideration of this work.', spacing_after=12)

    add_para(doc, 'Sincerely,', spacing_after=6)
    add_para(doc, '')
    add_para(doc, '[Corresponding Author Name]', bold=True)
    add_para(doc, '[Affiliation]')
    add_para(doc, '[Email Address]')
    add_para(doc, '[ORCID]')

    out_path = OUTPUT_DIR / 'cover_letter_english.docx'
    doc.save(str(out_path))
    print(f"English cover letter saved: {out_path}")
    return out_path


def create_japanese_cover_letter():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Date
    add_para(doc, '[日付]', font_size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Addressee
    add_para(doc, '', spacing_after=6)
    add_para(doc, 'The Editor', bold=True)
    add_para(doc, 'Scientific Data')
    add_para(doc, 'Nature Portfolio')
    add_para(doc, '')

    # Subject
    add_para(doc, 'Re: Article投稿 — "The Hidden Cost of Reproducibility: Commercial Software Dependency in Published Research and the Version Accessibility Gap"',
             bold=True, spacing_after=12)

    # Salutation
    add_para(doc, 'Dear Editor,', spacing_after=12)

    # Paragraph 1
    add_para(doc, f'このたび、"The Hidden Cost of Reproducibility: Commercial Software Dependency in Published Research and the Version Accessibility Gap"と題する論文を、Scientific Data誌のArticleとしてご査読いただきたく投稿いたします。本研究は、再現性の危機において十分に探究されていない重要な側面、すなわち特定バージョンが追試のために入手困難な商用ソフトウェアへの依存について取り組んでいます。',
             spacing_after=10)

    # Paragraph 2
    add_para(doc, f'PubMedの層別ランダムサンプリングにより、7つの研究分野（基礎医学、臨床医学、化学・材料科学、物理学・工学、社会・行動科学、計算科学、環境・地球科学）から2020年から2025年に出版された{N:,}本の論文を対象とした大規模横断研究を実施しました。95以上の正規表現パターンを用いて、各論文のソフトウェア言及、バージョン番号、ライセンス種別、推定追試コストを抽出しました。さらに、主要商用ソフトウェアベンダーの旧バージョンアクセスポリシーの体系的調査を実施し、実証データを補完しました。',
             spacing_after=10)

    # Paragraph 3
    add_para(doc, f'主要な知見は以下の通りです：(1) {overall["papers_with_commercial_sw"]/N*100:.1f}%の論文が商用ソフトウェアに依存しており、平均追試コストは1論文あたり${overall["mean_replication_cost_usd"]:,.0f}と推定されました。(2) 検出されたソフトウェアのうちバージョン番号が付随していたのはわずか{overall["mean_version_mention_rate"]*100:.1f}%でした。(3) 引用された商用ソフトウェアバージョンの97.0%が「入手困難」に分類されました。(4) 検証目的の「再現性ライセンス」を明示的に提供しているベンダーは皆無でした。我々は、出版された研究で引用されたソフトウェアバージョンと追試のために実際に入手可能なバージョンとの乖離を「バージョンアクセシビリティギャップ」として特定し、定量化可能かつ政策的に重要な再現性の障壁として提示します。',
             spacing_after=10)

    # Paragraph 4
    add_para(doc, '本論文がScientific DataのArticle形式に適する理由は以下の通りです。第一に、データ共有ポリシーと再現性インフラストラクチャという本誌のスコープの中核的トピックに直接取り組んでいます。第二に、{0:,}本の論文から抽出した35変数の完全なデータセットと、サンプリング・抽出パイプラインの全コードを公開しており、計算再現性に関するさらなる研究を可能にします。第三に、出版社、助成機関、ソフトウェアベンダーに対する即時的な政策的含意を有し、「再現性ライセンス」フレームワークを具体的に提案しています。7つの主要研究分野にまたがる学際的サンプルにより、本誌の多様な読者層に広く関連する内容となっています。'.format(N),
             spacing_after=10)

    # Paragraph 5
    add_para(doc, '本論文は他誌に投稿中または掲載済みではなく、全著者が投稿を承認しています。利益相反はありません。完全なデータセットと分析コードは、採択後に公開リポジトリに登録いたします。',
             spacing_after=10)

    # Paragraph 6
    add_para(doc, '査読者の推薦が必要な場合は、研究再現性、科学計量学、ソフトウェア持続可能性の専門家をご提案いたします。',
             spacing_after=10)

    # Closing
    add_para(doc, 'ご査読のほど、よろしくお願い申し上げます。', spacing_after=12)

    add_para(doc, '敬具', spacing_after=6)
    add_para(doc, '')
    add_para(doc, '[責任著者氏名]', bold=True)
    add_para(doc, '[所属機関]')
    add_para(doc, '[メールアドレス]')
    add_para(doc, '[ORCID]')

    out_path = OUTPUT_DIR / 'cover_letter_japanese.docx'
    doc.save(str(out_path))
    print(f"Japanese cover letter saved: {out_path}")
    return out_path


if __name__ == '__main__':
    create_english_cover_letter()
    create_japanese_cover_letter()
    print("\nBoth cover letters created successfully!")
